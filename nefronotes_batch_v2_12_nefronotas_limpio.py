# -*- coding: utf-8 -*-
"""
NefroNotes Batch Generator v2.11.0 - Nefro-Notas language sync
Sistema de generación masiva de notas nefrológicas mensuales
Desarrollado para Dr. Josué Tapia López - CMN Bajío IMSS

v2.11.0 - SINCRONIZACION CON GPT NEFRO-NOTAS:
       - Prompt maestro actualizado con reglas del GPT Nefro-Notas
       - Lenguaje medico natural, legalmente prudente y compacto
       - No inferir datos: solo usar informacion explicita
       - Promedios de SV y GID documentados sin mencionar numero de sesiones
       - Reglas de procedencia: IMSS indica; Hospital Cercano/Privados sugiere
       - Farmacos nuevos y calcimimeticos con lenguaje obligatorio "se sugiere"
       - Regla critica: no disminuir K del dializado; unidad usa 2K
       - DOCX Century Gothic 10, interlineado 1.5, carta, margenes indicados

v2.12.0 - PROMPT NEFRO-NOTAS DEPURADO:
       - Producto Ca x P calculado automaticamente cuando hay Ca y P en labs
       - Meta KDIGO CaxP <55 mg2/dL2 (riesgo calcificacion vascular si >=55)
       - Fecha de nacimiento (FN) extraida de notas HD via regex
       - Cache persistente de FN por expediente (pacientes_fn.json)
       - Edad calculada matematicamente desde FN a fecha de nota
       - Validacion post-generacion: corrige FN/Edad si Claude se equivoca
       - Nueva pestana "Fechas de Nacimiento" para gestion manual del cache
v2.9.1 - RANGOS KDIGO PARA INTERPRETACIÓN DE LABORATORIOS:
       - Rangos de electrolitos (K 3.5-5.5, Ca 8.4-10, P hasta 5)
       - Rangos de PTH en HD (150-300, evitar <100 o >500)
       - Rangos de calcidiol, ferritina, saturación transferrina
       - Regla crítica: NO diagnosticar falsos (K 3.8 no es hipopotasemia)
v2.9.0 - MEMORIA DE PROGRESO PERSISTENTE:
       - Si se corta la conexión, el progreso se guarda a disco
       - Al reiniciar, continúa donde quedó automáticamente
       - Las notas NO se borran hasta que tú lo digas
       - Botón explícito para limpiar progreso
v2.8.9 - Correcciones de formato y laboratorios:
       - FORZAR inclusión de sección "Estudios de laboratorio:" cuando hay CSV
       - Formato más compacto (menos saltos de línea)
       - Límite 4000 caracteres de TEXTO (no espacios)
       - Labs con formato abreviado obligatorio
v2.8.8 - Match de laboratorios por nombre (PatName)
v2.8.7 - Auto-guardado y continuación
v2.8.6 - Ajuste inteligente si nota >4000 chars
v2.8.5 - Correcciones de formato
v2.8.4 - Antecedentes relevantes, sin síntomas negativos
v2.8.3 - Abreviaturas obligatorias en laboratorios
v2.8.1 - Selector pacientes: EXPEDIENTE + NOMBRE
v2.8.0 - Fechas diagnósticos, firma única
"""

import streamlit as st
import pandas as pd
import anthropic
from pathlib import Path
import io
import zipfile
from datetime import datetime, date
import PyPDF2
import math
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import time
import json
from collections import Counter
import hashlib
import os

# ============================================================================
# v2.9.0: DIRECTORIO PARA PERSISTENCIA DE PROGRESO
# ============================================================================
PROGRESS_DIR = Path("NefroNotes_Progress")
PROGRESS_DIR.mkdir(parents=True, exist_ok=True)

PROGRESS_FILE = PROGRESS_DIR / "_progreso.json"
NOTES_FILE = PROGRESS_DIR / "_notas_generadas.json"

# v2.10.0: Cache persistente de fechas de nacimiento
FN_CACHE_FILE = PROGRESS_DIR / "pacientes_fn.json"

# ============================================================================
# v2.9.0: FUNCIONES DE PERSISTENCIA A DISCO
# ============================================================================

def cargar_progreso_de_disco():
    """Carga el progreso guardado desde disco"""
    progreso = {
        'completados': set(),
        'notas': [],
        'errores': [],
        'warnings': [],
        'ajustadas': []
    }
    
    # Cargar lista de completados
    if PROGRESS_FILE.exists():
        try:
            with open(PROGRESS_FILE, 'r', encoding='utf-8') as f:
                data = json.load(f)
            progreso['completados'] = set(data.get('completados', []))
            progreso['errores'] = data.get('errores', [])
            progreso['warnings'] = data.get('warnings', [])
            progreso['ajustadas'] = data.get('ajustadas', [])
        except Exception as e:
            st.warning(f"⚠️ Error cargando progreso: {e}")
    
    # Cargar notas generadas
    if NOTES_FILE.exists():
        try:
            with open(NOTES_FILE, 'r', encoding='utf-8') as f:
                data = json.load(f)
            progreso['notas'] = data.get('notas', [])
        except Exception as e:
            st.warning(f"⚠️ Error cargando notas: {e}")
    
    return progreso

def guardar_nota_a_disco(nota):
    """Guarda una nota individual inmediatamente después de generarla"""
    try:
        # Cargar notas existentes
        notas = []
        if NOTES_FILE.exists():
            with open(NOTES_FILE, 'r', encoding='utf-8') as f:
                data = json.load(f)
            notas = data.get('notas', [])
        
        # Agregar nueva nota (doc_bytes en base64)
        import base64
        nota_serializable = {
            'paciente': nota['paciente'],
            'expediente': nota['expediente'],
            'content': nota['content'],
            'doc_bytes_b64': base64.b64encode(nota['doc_bytes']).decode('utf-8'),
            'tiene_labs': nota.get('tiene_labs', False),
            'tiene_hd_notes': nota.get('tiene_hd_notes', False),
            'fue_ajustada': nota.get('fue_ajustada', False),
            'timestamp': datetime.now().isoformat()
        }
        notas.append(nota_serializable)
        
        # Guardar a disco
        with open(NOTES_FILE, 'w', encoding='utf-8') as f:
            json.dump({'notas': notas, 'ultima_actualizacion': datetime.now().isoformat()}, f)
        
        return True
    except Exception as e:
        st.error(f"❌ Error guardando nota a disco: {e}")
        return False

def guardar_progreso_a_disco(expediente_completado, errores=None, warnings=None, ajustadas=None):
    """Actualiza el archivo de progreso con un expediente completado"""
    try:
        # Cargar progreso existente
        data = {
            'completados': [],
            'errores': [],
            'warnings': [],
            'ajustadas': []
        }
        if PROGRESS_FILE.exists():
            with open(PROGRESS_FILE, 'r', encoding='utf-8') as f:
                data = json.load(f)
        
        # Actualizar
        if expediente_completado and expediente_completado not in data['completados']:
            data['completados'].append(expediente_completado)
        if errores:
            data['errores'] = errores
        if warnings:
            data['warnings'] = warnings
        if ajustadas:
            data['ajustadas'] = ajustadas
        
        data['ultima_actualizacion'] = datetime.now().isoformat()
        data['total_completados'] = len(data['completados'])
        
        # Guardar
        with open(PROGRESS_FILE, 'w', encoding='utf-8') as f:
            json.dump(data, f)
        
        return True
    except Exception as e:
        st.error(f"❌ Error guardando progreso: {e}")
        return False

def limpiar_progreso_de_disco():
    """Elimina todos los archivos de progreso"""
    try:
        if PROGRESS_FILE.exists():
            PROGRESS_FILE.unlink()
        if NOTES_FILE.exists():
            NOTES_FILE.unlink()
        return True
    except Exception as e:
        st.error(f"❌ Error limpiando progreso: {e}")
        return False

def hay_progreso_pendiente():
    """Verifica si hay progreso guardado en disco"""
    return PROGRESS_FILE.exists() or NOTES_FILE.exists()

def obtener_resumen_progreso():
    """Obtiene un resumen del progreso guardado"""
    if not hay_progreso_pendiente():
        return None
    
    resumen = {
        'completados': 0,
        'notas': 0,
        'errores': 0,
        'ultima_actualizacion': None
    }
    
    if PROGRESS_FILE.exists():
        try:
            with open(PROGRESS_FILE, 'r', encoding='utf-8') as f:
                data = json.load(f)
            resumen['completados'] = len(data.get('completados', []))
            resumen['errores'] = len(data.get('errores', []))
            resumen['ultima_actualizacion'] = data.get('ultima_actualizacion')
        except:
            pass
    
    if NOTES_FILE.exists():
        try:
            with open(NOTES_FILE, 'r', encoding='utf-8') as f:
                data = json.load(f)
            resumen['notas'] = len(data.get('notas', []))
        except:
            pass
    
    return resumen

def cargar_notas_de_disco():
    """Carga las notas completas desde disco para descarga"""
    notas = []
    if NOTES_FILE.exists():
        try:
            import base64
            with open(NOTES_FILE, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            for nota in data.get('notas', []):
                nota_completa = {
                    'paciente': nota['paciente'],
                    'expediente': nota['expediente'],
                    'content': nota['content'],
                    'doc_bytes': base64.b64decode(nota['doc_bytes_b64']),
                    'tiene_labs': nota.get('tiene_labs', False),
                    'tiene_hd_notes': nota.get('tiene_hd_notes', False),
                    'fue_ajustada': nota.get('fue_ajustada', False)
                }
                notas.append(nota_completa)
        except Exception as e:
            st.error(f"❌ Error cargando notas: {e}")
    
    return notas

# ============================================================================
# v2.10.0: GESTIÓN DE FECHAS DE NACIMIENTO (FN)
# ============================================================================

def cargar_cache_fn():
    """Carga el caché de fechas de nacimiento desde disco.
    Retorna dict {expediente: 'dd/mm/aaaa'}"""
    if FN_CACHE_FILE.exists():
        try:
            with open(FN_CACHE_FILE, 'r', encoding='utf-8') as f:
                data = json.load(f)
            return data.get('pacientes', {})
        except Exception as e:
            st.warning(f"⚠️ Error cargando caché de FN: {e}")
    return {}

def guardar_cache_fn(cache_dict):
    """Guarda el caché completo de fechas de nacimiento a disco."""
    try:
        data = {
            'pacientes': cache_dict,
            'ultima_actualizacion': datetime.now().isoformat(),
            'total': len(cache_dict)
        }
        with open(FN_CACHE_FILE, 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
        return True
    except Exception as e:
        st.error(f"❌ Error guardando caché de FN: {e}")
        return False

def agregar_fn_a_cache(expediente, fecha_nacimiento_str):
    """Agrega o actualiza una FN al caché. Acepta 'dd/mm/aaaa'."""
    if not expediente or not fecha_nacimiento_str:
        return False
    
    # Normalizar expediente
    expediente = str(expediente).strip().upper()
    
    # Validar formato dd/mm/aaaa
    try:
        datetime.strptime(fecha_nacimiento_str, '%d/%m/%Y')
    except ValueError:
        return False
    
    cache = cargar_cache_fn()
    cache[expediente] = fecha_nacimiento_str
    return guardar_cache_fn(cache)

def extraer_fecha_nacimiento_de_notas(hd_notes_text):
    """Extrae la fecha de nacimiento (FN) de las notas HD mediante regex.
    Busca patrones como 'Fecha de nacimiento: dd/mm/aaaa' o 'Fecha Nacim. dd/mm/aaaa'.
    Retorna 'dd/mm/aaaa' o None si no encuentra."""
    if not hd_notes_text:
        return None
    
    # Patrones de búsqueda en orden de preferencia
    patrones = [
        r'[Ff]echa\s+de\s+[Nn]acimiento\s*:?\s*(\d{1,2})[/\-\.](\d{1,2})[/\-\.](\d{4})',
        r'[Ff]echa\s+[Nn]acim\.?\s*:?\s*(\d{1,2})[/\-\.](\d{1,2})[/\-\.](\d{4})',
        r'[Ff]\.?\s*[Nn]acimiento\s*:?\s*(\d{1,2})[/\-\.](\d{1,2})[/\-\.](\d{4})',
        r'[Nn]ac\.?\s*:?\s*(\d{1,2})[/\-\.](\d{1,2})[/\-\.](\d{4})',
    ]
    
    # Recopilar todas las fechas encontradas
    fechas_encontradas = []
    for patron in patrones:
        matches = re.findall(patron, hd_notes_text)
        for match in matches:
            dia, mes, anio = match
            try:
                fecha_obj = datetime(int(anio), int(mes), int(dia))
                # Validación: FN razonable (entre 1900 y hace 18 años)
                if 1900 <= fecha_obj.year <= datetime.now().year - 18:
                    fechas_encontradas.append(fecha_obj.strftime('%d/%m/%Y'))
            except ValueError:
                continue
    
    if not fechas_encontradas:
        return None
    
    # Si hay varias, devolver la más frecuente
    contador = Counter(fechas_encontradas)
    return contador.most_common(1)[0][0]

def calcular_edad(fecha_nacimiento_str, fecha_referencia=None):
    """Calcula la edad correctamente desde FN a fecha de referencia.
    Considera si el cumpleaños ya pasó en el año actual.
    
    Args:
        fecha_nacimiento_str: 'dd/mm/aaaa'
        fecha_referencia: date object o None (usa hoy)
    
    Retorna: edad en años (int) o None si error
    """
    if not fecha_nacimiento_str:
        return None
    
    try:
        fn = datetime.strptime(fecha_nacimiento_str, '%d/%m/%Y').date()
    except (ValueError, TypeError):
        return None
    
    if fecha_referencia is None:
        fecha_referencia = date.today()
    elif isinstance(fecha_referencia, datetime):
        fecha_referencia = fecha_referencia.date()
    
    edad = fecha_referencia.year - fn.year
    # Si el cumpleaños aún no ha pasado este año, restar 1
    if (fecha_referencia.month, fecha_referencia.day) < (fn.month, fn.day):
        edad -= 1
    
    return edad

def obtener_fn_y_edad(expediente, hd_notes_text, fecha_nota):
    """Orquesta la obtención de FN y cálculo de edad.
    Prioridad: 1) Caché 2) Extracción de notas HD (y auto-guarda al caché)
    
    Retorna tupla (fn_str, edad, origen) donde origen es 'cache', 'notas', o 'no_encontrado'
    """
    expediente = str(expediente).strip().upper()
    
    # 1. Buscar en caché
    cache = cargar_cache_fn()
    if expediente in cache:
        fn_str = cache[expediente]
        edad = calcular_edad(fn_str, fecha_nota)
        return fn_str, edad, 'cache'
    
    # 2. Extraer de notas HD
    fn_extraida = extraer_fecha_nacimiento_de_notas(hd_notes_text)
    if fn_extraida:
        # Auto-guardar al caché para uso futuro
        agregar_fn_a_cache(expediente, fn_extraida)
        edad = calcular_edad(fn_extraida, fecha_nota)
        return fn_extraida, edad, 'notas'
    
    # 3. No encontrado
    return None, None, 'no_encontrado'

def validar_y_corregir_fn_edad_en_nota(note_content, fn_correcta, edad_correcta):
    """Valida que la nota generada tenga la FN y edad correctas.
    Si Claude escribió valores distintos, los corrige.
    
    Retorna tupla (nota_corregida, fue_corregida_bool)
    """
    if not fn_correcta or edad_correcta is None:
        return note_content, False
    
    fue_corregida = False
    nota = note_content
    
    # Corregir línea "Fecha de nacimiento: ..."
    # Importante: usar [ \t]* (no \s*) para no cruzar saltos de línea
    patron_fn = r'(Fecha de nacimiento:[ \t]*)([^\n]+)'
    match_fn = re.search(patron_fn, nota)
    if match_fn:
        fn_escrita = match_fn.group(2).strip()
        if fn_escrita != fn_correcta:
            nota = re.sub(patron_fn, f'\\g<1>{fn_correcta}', nota, count=1)
            fue_corregida = True
    
    # Corregir línea "Edad: X años"
    patron_edad = r'(Edad:[ \t]*)(\d+\s*años?)'
    match_edad = re.search(patron_edad, nota)
    if match_edad:
        edad_escrita_str = match_edad.group(2)
        edad_escrita_num = int(re.search(r'\d+', edad_escrita_str).group())
        if edad_escrita_num != edad_correcta:
            nota = re.sub(patron_edad, f'\\g<1>{edad_correcta} años', nota, count=1)
            fue_corregida = True
    
    # También corregir referencia a edad en Diagnósticos: "Masculino de X años" o "Femenina de X años"
    patron_diag = r'(Masculino|Femenina|Femenino|Masculina)\s+de\s+(\d+)\s+años'
    match_diag = re.search(patron_diag, nota)
    if match_diag:
        edad_diag = int(match_diag.group(2))
        if edad_diag != edad_correcta:
            nota = re.sub(
                patron_diag,
                lambda m: f"{m.group(1)} de {edad_correcta} años",
                nota,
                count=1
            )
            fue_corregida = True
    
    return nota, fue_corregida

# ============================================================================
# SISTEMA DE AUTENTICACIÓN
# ============================================================================

def check_password():
    """
    Verifica la contraseña de acceso a la aplicación.
    La contraseña se configura en secrets.toml o se usa una por defecto.
    """
    default_password = "NefroNotes2025"
    
    try:
        if "passwords" in st.secrets and "admin" in st.secrets["passwords"]:
            target_password = st.secrets["passwords"]["admin"]
        else:
            target_password = default_password
    except Exception:
        target_password = default_password
    
    def password_entered():
        entered_password = st.session_state.get("password", "")
        if entered_password == target_password:
            st.session_state["password_correct"] = True
            if "password" in st.session_state:
                del st.session_state["password"]
        else:
            st.session_state["password_correct"] = False

    if st.session_state.get("password_correct", False):
        return True

    st.markdown("""
        <style>
        .login-container {
            max-width: 400px;
            margin: 100px auto;
            padding: 40px;
            background: white;
            border-radius: 10px;
            box-shadow: 0 4px 6px rgba(0,0,0,0.1);
        }
        </style>
    """, unsafe_allow_html=True)
    
    col1, col2, col3 = st.columns([1, 2, 1])
    
    with col2:
        st.markdown("# 🔐")
        st.markdown("## NefroNotes")
        st.markdown("### Acceso Privado")
        st.markdown("---")
        
        st.text_input(
            "Contraseña de Acceso",
            type="password",
            on_change=password_entered,
            key="password",
            placeholder="Ingresa tu contraseña"
        )
        
        if "password_correct" in st.session_state and not st.session_state["password_correct"]:
            st.error("🚫 Contraseña incorrecta. Inténtalo de nuevo.")
        
        st.markdown("---")
        st.caption("💡 **Contraseña por defecto:** NefroNotes2025")
        st.caption("📝 Para cambiarla, configura `secrets.toml`")
        st.caption("🏥 Dr. Josué Tapia López - CMN Bajío IMSS")
    
    return False

# ============================================================================
# VERIFICAR AUTENTICACIÓN
# ============================================================================

if not check_password():
    st.stop()

# ============================================================================

# Mapeo de meses en español
MESES_ESP = {
    1: 'enero', 2: 'febrero', 3: 'marzo', 4: 'abril',
    5: 'mayo', 6: 'junio', 7: 'julio', 8: 'agosto',
    9: 'septiembre', 10: 'octubre', 11: 'noviembre', 12: 'diciembre'
}

def detectar_mes_de_notas(hd_notes_text):
    """Detecta el mes predominante en las notas de hemodiálisis"""
    if not hd_notes_text:
        mes_num = datetime.now().month
        return MESES_ESP[mes_num]
    
    fechas_encontradas = []
    patron_fecha = r'\b(\d{1,2})[/\-\.](\d{1,2})[/\-\.](\d{4})\b'
    matches = re.findall(patron_fecha, hd_notes_text)
    
    for match in matches:
        try:
            dia, mes, anio = match
            mes_num = int(mes)
            if 1 <= mes_num <= 12:
                fechas_encontradas.append(mes_num)
        except:
            continue
    
    if fechas_encontradas:
        mes_comun = Counter(fechas_encontradas).most_common(1)[0][0]
        return MESES_ESP[mes_comun]
    
    mes_num = datetime.now().month
    return MESES_ESP[mes_num]

def detectar_tipo_expediente(expediente):
    """Detecta el tipo de expediente basado en el código"""
    if not expediente:
        return "OTRO"
    
    expediente_upper = str(expediente).upper().strip()
    
    if expediente_upper.startswith('LE'):
        return "LE"
    elif expediente_upper.startswith('RL'):
        return "RL"
    elif expediente_upper.startswith('BR'):
        return "BR"
    else:
        return "OTRO"

# Configuración de la página
st.set_page_config(
    page_title="NefroNotes Batch Generator",
    page_icon="🩺",
    layout="wide"
)

# CSS personalizado
st.markdown("""
<style>
    .main-header {
        font-size: 2.5rem;
        color: #1E40AF;
        font-weight: bold;
        text-align: center;
        margin-bottom: 2rem;
    }
    .step-header {
        background: linear-gradient(90deg, #3B82F6 0%, #1E40AF 100%);
        color: white;
        padding: 1rem;
        border-radius: 0.5rem;
        margin: 1rem 0;
    }
    .success-box {
        background-color: #D1FAE5;
        border-left: 4px solid #10B981;
        padding: 1rem;
        border-radius: 0.5rem;
    }
    .warning-box {
        background-color: #FEF3C7;
        border-left: 4px solid #F59E0B;
        padding: 1rem;
        border-radius: 0.5rem;
    }
    .progress-box {
        background-color: #DBEAFE;
        border-left: 4px solid #3B82F6;
        padding: 1rem;
        border-radius: 0.5rem;
        margin: 0.5rem 0;
    }
    .disk-saved-box {
        background-color: #FEF3C7;
        border-left: 4px solid #F59E0B;
        padding: 1rem;
        border-radius: 0.5rem;
        margin: 0.5rem 0;
    }
</style>
""", unsafe_allow_html=True)

# ============================================================================
# FUNCIONES AUXILIARES
# ============================================================================

def parse_hemoHL7_consolidated_pdf(pdf_file):
    """Parsea un PDF consolidado de laboratorios con múltiples pacientes."""
    try:
        pdf_file.seek(0)
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        full_text = ""
        for page in pdf_reader.pages:
            full_text += page.extract_text() + "\n"
        
        labs_by_expediente = {}
        
        pattern_exped = r'Exped:[\s\S]{0,200}?([A-Z]{2}\d{4})'
        pattern_paciente = r'Paciente:\s*([A-Z]{2}\d{4})'
        
        matches_exped = list(re.finditer(pattern_exped, full_text))
        matches_paciente = list(re.finditer(pattern_paciente, full_text))
        
        all_matches = []
        
        for match in matches_exped:
            all_matches.append({
                'pos': match.start(),
                'expediente': match.group(1).strip().upper(),
                'match': match
            })
        
        for match in matches_paciente:
            all_matches.append({
                'pos': match.start(),
                'expediente': match.group(1).strip().upper(),
                'match': match
            })
        
        all_matches.sort(key=lambda x: x['pos'])
        
        for i, match_info in enumerate(all_matches):
            expediente = match_info['expediente']
            start_pos = match_info['pos']
            if i + 1 < len(all_matches):
                end_pos = all_matches[i + 1]['pos']
            else:
                end_pos = len(full_text)
            
            patient_text = full_text[start_pos:end_pos]
            labs_by_expediente[expediente] = patient_text
        
        return labs_by_expediente
        
    except Exception as e:
        st.error(f"Error al parsear PDF consolidado de laboratorios: {str(e)}")
        return {}

def normalize_censo_columns(df):
    """Normaliza las columnas del censo para que sean consistentes."""
    df_normalized = df.copy()
    
    expediente_col = None
    possible_expediente_names = ['exped', 'expediente', 'exp', 'num', 'numero', 'no.', 'id']
    
    for col in df_normalized.columns:
        col_lower = str(col).lower().strip()
        if any(name in col_lower for name in possible_expediente_names):
            expediente_col = col
            break
    
    if expediente_col is None:
        for col in df_normalized.columns:
            sample = df_normalized[col].astype(str).head(3)
            if any(re.match(r'[A-Z]{2}\d+', str(val), re.IGNORECASE) for val in sample):
                expediente_col = col
                break
    
    if expediente_col is None:
        expediente_col = df_normalized.columns[0]
    
    nombre_col = None
    possible_nombre_names = ['paciente', 'nombre', 'apellido']
    
    for col in df_normalized.columns:
        if col == expediente_col:
            continue
        if str(col).startswith('Unnamed'):
            continue
        col_lower = str(col).lower().strip()
        if any(name in col_lower for name in possible_nombre_names):
            nombre_col = col
            break
    
    if nombre_col is None:
        for col in df_normalized.columns:
            if col != expediente_col:
                sample_values = df_normalized[col].astype(str).str.strip()
                if sample_values.str.match(r'^\d+$').all():
                    continue
                nombre_col = col
                break
    
    if nombre_col is None:
        for col in df_normalized.columns:
            if col != expediente_col:
                nombre_col = col
                break
    
    column_mapping = {
        expediente_col: 'Exped.2',
        nombre_col: 'Paciente'
    }
    
    for col in df_normalized.columns:
        col_lower = str(col).lower().strip()
        if 'proced' in col_lower or 'procedencia' in col_lower:
            column_mapping[col] = 'PROCED'
        elif 'sala' in col_lower:
            column_mapping[col] = 'SALA'
        elif 'turno' in col_lower or 'horario' in col_lower:
            column_mapping[col] = 'TURNO'
    
    df_normalized = df_normalized.rename(columns=column_mapping)
    
    if 'PROCED' not in df_normalized.columns:
        df_normalized['PROCED'] = 'N/A'
    if 'SALA' not in df_normalized.columns:
        df_normalized['SALA'] = 'N/A'
    if 'TURNO' not in df_normalized.columns:
        df_normalized['TURNO'] = 'N/A'
    
    df_final = df_normalized[['Exped.2', 'Paciente', 'PROCED', 'SALA', 'TURNO']].copy()
    df_final['Exped.2'] = df_final['Exped.2'].astype(str).str.strip().str.upper()
    df_final['Paciente'] = df_final['Paciente'].astype(str).str.strip()
    
    return df_final

def extract_text_from_pdf(pdf_file):
    """Extrae texto de un archivo PDF"""
    try:
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
        return text
    except Exception as e:
        return f"Error al leer PDF: {str(e)}"

def normalizar_nombre(nombre):
    """Normaliza un nombre para comparación (quita acentos, comas, puntos, mayúsculas, espacios extra)"""
    if not nombre:
        return ""
    import unicodedata
    # Quitar acentos
    nombre = unicodedata.normalize('NFD', str(nombre))
    nombre = ''.join(c for c in nombre if unicodedata.category(c) != 'Mn')
    # Quitar comas (diferencia entre formato censo y CSV)
    nombre = nombre.replace(',', '')
    # Quitar puntos (MA. vs MA, J. vs J)
    nombre = nombre.replace('.', '')
    # Mayúsculas y quitar espacios extra
    nombre = ' '.join(nombre.upper().split())
    return nombre

def process_labs_data(labs_df, expediente, nombre_paciente=None, hd_notes_text=None, labs_pdf_text=None, hemoHL7_text=None):
    """Procesa los datos de laboratorio para un paciente específico.
    Busca primero por ExpedientNum, si no encuentra busca por nombre (PatName)."""
    try:
        if labs_df is not None:
            patient_labs = None
            
            # 1. Intentar buscar por ExpedientNum
            if 'ExpedientNum' in labs_df.columns:
                patient_labs = labs_df[labs_df['ExpedientNum'] == expediente]
                
                if patient_labs.empty and expediente:
                    exp_numeros = ''.join(filter(str.isdigit, str(expediente)))
                    if exp_numeros:
                        patient_labs = labs_df[labs_df['ExpedientNum'].astype(str).str.contains(exp_numeros, na=False)]
            
            # 2. Si no encontró por expediente, buscar por nombre
            if (patient_labs is None or patient_labs.empty) and nombre_paciente and 'PatName' in labs_df.columns:
                nombre_normalizado = normalizar_nombre(nombre_paciente)
                
                labs_df['_nombre_norm'] = labs_df['PatName'].apply(normalizar_nombre)
                
                # Buscar coincidencia exacta
                patient_labs = labs_df[labs_df['_nombre_norm'] == nombre_normalizado]
                
                # Si no hay coincidencia exacta, buscar por apellidos
                if patient_labs.empty:
                    apellidos_paciente = tuple(nombre_normalizado.split()[:2])
                    for idx, row in labs_df.iterrows():
                        apellidos_csv = tuple(row['_nombre_norm'].split()[:2])
                        if apellidos_paciente == apellidos_csv:
                            patient_labs = labs_df.iloc[[idx]]
                            break
                
                if '_nombre_norm' in labs_df.columns:
                    labs_df.drop('_nombre_norm', axis=1, inplace=True)
            
            # 3. Verificar si encontró y si tiene valores válidos (no NaN)
            if patient_labs is not None and not patient_labs.empty:
                latest_lab = patient_labs.iloc[0]
                
                # Verificar si tiene al menos Hgb válido (indicador de que tiene labs)
                hgb_val = latest_lab.get('Hgb') if 'Hgb' in latest_lab.index else None
                
                import math
                if hgb_val is None or (isinstance(hgb_val, float) and math.isnan(hgb_val)):
                    # Paciente está en CSV pero sin valores → extraer de notas
                    if hemoHL7_text:
                        return None, "EXTRACT_FROM_HEMOHL7"
                    elif labs_pdf_text:
                        return None, "EXTRACT_FROM_PDF_LABS"
                    elif hd_notes_text:
                        return None, "EXTRACT_FROM_NOTES"
                    else:
                        return None, "NO_LABS"
                
                # Tiene valores válidos - construir texto de labs en FORMATO COMPACTO
                def safe_val(val):
                    """Devuelve el valor o vacío si es NaN"""
                    if val is None:
                        return None
                    if isinstance(val, float) and math.isnan(val):
                        return None
                    return val
                
                fecha = safe_val(latest_lab.get('DrawDate', ''))
                
                # Construir labs en formato compacto (estilo Dr. Tapia)
                labs_parts = []
                
                # Biometría
                if safe_val(latest_lab.get('Hgb')): labs_parts.append(f"Hb {latest_lab['Hgb']}")
                if safe_val(latest_lab.get('Hct')): labs_parts.append(f"Hto {latest_lab['Hct']}%")
                if safe_val(latest_lab.get('WBC')): labs_parts.append(f"Leu {latest_lab['WBC']}")
                if safe_val(latest_lab.get('PlateletCt')): labs_parts.append(f"PLQ {latest_lab['PlateletCt']}")
                
                # Química
                if safe_val(latest_lab.get('Glucose')): labs_parts.append(f"Glu {latest_lab['Glucose']}")
                if safe_val(latest_lab.get('Creatinine')): labs_parts.append(f"Cr {latest_lab['Creatinine']}")
                if safe_val(latest_lab.get('BUN_Pre')): labs_parts.append(f"BUN {latest_lab['BUN_Pre']}")
                if safe_val(latest_lab.get('URR')): labs_parts.append(f"URR {latest_lab['URR']}%")
                
                # Electrolitos
                if safe_val(latest_lab.get('Sodium')): labs_parts.append(f"Na {latest_lab['Sodium']}")
                if safe_val(latest_lab.get('Potassium')): labs_parts.append(f"K {latest_lab['Potassium']}")
                
                # v2.10.0: Ca, P y producto CaxP calculado automáticamente
                ca_val = safe_val(latest_lab.get('Calcium'))
                p_val = safe_val(latest_lab.get('Phosphorus'))
                if ca_val is not None:
                    labs_parts.append(f"Ca {ca_val}")
                if p_val is not None:
                    labs_parts.append(f"P {p_val}")
                # Calcular CaxP si ambos están disponibles
                if ca_val is not None and p_val is not None:
                    try:
                        caxp = round(float(ca_val) * float(p_val), 1)
                        labs_parts.append(f"CaxP {caxp}")
                    except (ValueError, TypeError):
                        pass
                
                # Otros
                if safe_val(latest_lab.get('Albumin')): labs_parts.append(f"Alb {latest_lab['Albumin']}")
                if safe_val(latest_lab.get('Ferritin')): labs_parts.append(f"Ferritina {latest_lab['Ferritin']}")
                if safe_val(latest_lab.get('PTHIntact')): labs_parts.append(f"PTH {latest_lab['PTHIntact']}")
                
                labs_compacto = ", ".join(labs_parts)
                
                lab_text = f"""
LABORATORIOS DEL CSV (Fecha: {fecha}):
{labs_compacto}

FORMATO PARA LA NOTA:
Estudios de laboratorio: {fecha}: {labs_compacto}.
"""
                return lab_text, "OK"
            
            # No encontró en CSV
            if hemoHL7_text:
                return None, "EXTRACT_FROM_HEMOHL7"
            elif labs_pdf_text:
                return None, "EXTRACT_FROM_PDF_LABS"
            elif hd_notes_text:
                return None, "EXTRACT_FROM_NOTES"
            else:
                return None, "NO_LABS"
        
        elif hemoHL7_text:
            return None, "EXTRACT_FROM_HEMOHL7"
        
        elif labs_pdf_text:
            return None, "EXTRACT_FROM_PDF_LABS"
        
        elif hd_notes_text:
            return None, "EXTRACT_FROM_NOTES"
        
        else:
            return None, "NO_LABS"
            
    except Exception as e:
        return None, f"ERROR: {str(e)}"


# ============================================================================
# v2.11.0: PROMPT MAESTRO SINCRONIZADO CON GPT "NEFRO-NOTAS"
# ============================================================================

NEFRONOTAS_MASTER_RULES = """
═══════════════════════════════════════════════════════════════════════════════
PROMPT MAESTRO NEFRO-NOTAS v2.11.0 - REGLAS SUPERIORES
═══════════════════════════════════════════════════════════════════════════════

Estas reglas sustituyen cualquier instrucción previa que entre en conflicto. La nota debe parecer escrita por un nefrólogo con experiencia clínica, no por una IA: lenguaje natural, sobrio, técnico, compacto y legalmente prudente.

1. LÍMITES Y ESTILO FINAL:
- Máximo absoluto: 4000 caracteres, incluyendo espacios, puntuación y saltos de línea.
- Objetivo operativo: 3800-3900 caracteres para dejar margen de seguridad.
- Máximo 2 cuartillas al imprimir a PDF.
- Compactar secciones en párrafos corridos; evitar listados largos, frases repetitivas y lenguaje robótico.
- Texto plano, sin markdown, sin viñetas, sin referencias bibliográficas explícitas.

2. NO INFERIR DATOS:
- Usar solo información explícita en censo, notas diarias, nota mensual previa, PDFs o CSV de laboratorio.
- Las notas diarias son muestra representativa del mes, no registro completo: NO escribir ni sugerir que son todas las sesiones del mes.
- Notas diarias tienen prioridad sobre nota mensual previa para medicamentos, dosis, parámetros, acceso vascular y eventos recientes.
- Si un dato no está documentado, omitirlo o redactar de forma neutral. No inventar valores “razonables”.
- No usar frases tipo “no disponible/no especificado/no documentado” salvo indicación específica de laboratorios faltantes por tipo de expediente.

3. CÁLCULOS Y PROMEDIOS:
- Signos vitales: usar promedio aritmético de las sesiones documentadas cuando existan varios registros; si solo hay uno, usar ese registro.
- Ganancia interdialítica: reportar promedio y rango si los datos lo permiten, por ejemplo “GID promedio 2.1 kg, rango 1.7-3.1 kg”.
- NO mencionar el número de sesiones documentadas.
- URR: solo si hay BUN/urea pre y post o valor ya reportado. No inventar URR si no se puede calcular con datos explícitos.

4. ÁCIDO DIALIZADO Y POTASIO:
- La unidad usa SOLO concentración 2K.
- NO sugerir disminuir K en ácido dializado.
- NO indicar resinas tipo Lokelma/patiromer como plan rutinario; los pacientes usualmente no las compran.
- Para hiperpotasemia: reforzar dieta baja en K, restricción hídrica y adherencia a sesiones.
- Furosemida SOLO si hay uresis residual explícitamente documentada.

5. PROCEDENCIA, FRECUENCIA Y DURACIÓN:
- IMSS/HGR 58/IMSS Gto.: puede indicarse continuidad o ajuste de 3 sesiones/semana si está documentado y procede.
- Hospital Cercano: 2-3 sesiones/semana; si requiere más, redactar “se sugiere incrementar frecuencia a 3 sesiones/semana”, NO como indicación directa.
- Privados: 1-3 sesiones según capacidad económica; usar lenguaje de sugerencia para incrementos.
- Duración estándar 180 min; máximo práctico 210 min por limitaciones operativas. Si se requiere más por URR bajo, sobrecarga o hiperpotasemia, redactar “se sugiere incrementar duración de sesión a X minutos”.
- Estas notas pueden revisarse en auditorías, procesos administrativos o juzgados; usar lenguaje clínico prudente.

6. MEDICAMENTOS - LENGUAJE OBLIGATORIO:
- Cualquier medicamento nuevo: “se sugiere iniciar [medicamento] [dosis] [vía] [frecuencia]”.
- Calcimiméticos: “se sugiere uso de cinacalcet [dosis]”, no “iniciar cinacalcet”.
- Carbonato de calcio: escribir “carbonato de calcio 500 mg”; NO escribir “Tums”.
- Si hay calcio alto/alto-normal con hiperfosfatemia: “se sugiere uso de quelante no cálcico, sevelamero [dosis]”.
- No sugerir ajustes adicionales si ya fueron valorados o modificados en notas diarias recientes.
- Tratamiento crónico debe ir en línea corrida; no incluir EPO ni hierro IV ahí si se reportan en parámetros o plan dialítico.

7. SELLO DE CATÉTER:
- Solo mencionar sello si el acceso es catéter. NO usar sello para FAVI.
- Heparina estándar: escribir “heparina” sin especificar 1000 UI/ml, salvo que esté explícito.
- Heparina alta: escribir “heparina 5000 UI/ml” solo si está explícito.
- Citra-Lock 46.7%: usar cuando esté documentado y el catéter funcione bien.
- Heparina preferida si hay disfunción por fibrina/coágulos documentada.

8. ACCESO VASCULAR:
- Catéter tunelizado funcional: NO sugerir cambio a FAVI por rutina.
- Solo sugerir cambio o valoración de acceso si hay complicaciones documentadas: bacteriemia, disfunción persistente, trombosis, infección, mal flujo o agotamiento de acceso.

9. LABORATORIOS Y METAS:
- Usar los valores más recientes con fecha clara.
- Comparar contra metas KDIGO/KDOQI: Hb 10-12 g/dL, K >5.1 relevante como hiperpotasemia clínica según contexto, P >5.1 hiperfosfatemia, Ca 8.5-10.4, URR >65%.
- Interpretar sin sobrediagnosticar: no llamar hiperfosfatemia significativa a elevaciones mínimas si clínicamente no amerita ajuste.
- Si Ca y P están disponibles, considerar CaxP; riesgo relevante si ≥55.

10. PARÁMETROS DE DIÁLISIS:
- Parámetros en una sola línea corrida: sesiones/semana, filtro, duración, QB/QS, QD, peso seco, temperatura, heparina, K 2, HCO3 35, Ca 3.5, UF máxima 13 ml/kg/h si aplica, sello solo si catéter, EPO si está documentada.

11. ESTRUCTURA FINAL OBLIGATORIA:
Encabezado demográfico en líneas separadas: Nombre, Fecha de nacimiento si consta, Edad si consta, Inicio de hemodiálisis, Fecha de ingreso a Clínica, Expediente, Fecha, Hora.
Título centrado: Nota Nefrología - Nota mensual [mes].
Diagnósticos: párrafo corrido.
Evolución del mes: párrafo corrido con eventos positivos/relevantes; no listar síntomas negados.
Signos vitales y Exploración física: líneas compactas, preferentemente consecutivas.
Estudios de laboratorio: fecha en formato claro y valores en línea compacta.
Análisis: párrafo conciso integrando clínica, metas y problemas activos.
Plan: párrafo conciso con continuidad, ajustes y lenguaje “se sugiere” cuando corresponda.
Tratamiento crónico: línea corrida, sin EPO ni Fe IV si ya se reportan por diálisis.
Parámetros de diálisis: todo en línea corrida.
Pronóstico: línea corrida.
Firma: NO incluir; el sistema la agrega automáticamente.

12. VERIFICACIÓN FINAL:
Antes de entregar, revisar mentalmente: ≤4000 caracteres; medicamentos/dosis coinciden con notas diarias más recientes; no hay inferencias; no se sugiere disminuir K del dializado; no hay Lokelma; sello solo en catéter; no hay Tums; no hay cambio de catéter tunelizado funcional a FAVI sin justificación; lenguaje de sugerencia en Hospital Cercano/Privados y en medicamentos nuevos.
═══════════════════════════════════════════════════════════════════════════════
"""

def create_master_prompt(patient_data, labs_data, labs_status, hd_notes_text, has_hd_notes, labs_pdf_text="", hemoHL7_text="", fecha_nota=None, fn_paciente=None, edad_paciente=None):
    """Crea el prompt master para Claude con todos los datos del paciente.
    
    v2.10.0: fn_paciente y edad_paciente son obligatorios cuando se conocen.
    Si se proporcionan, Claude DEBE usar esos valores exactos."""
    
    if fecha_nota:
        mes_nota = MESES_ESP[fecha_nota.month]
        anio_nota = fecha_nota.year
        fecha_str = fecha_nota.strftime("%d/%m/%Y")
    else:
        mes_nota = detectar_mes_de_notas(hd_notes_text)
        anio_nota = datetime.now().year
        fecha_str = datetime.now().strftime("%d/%m/%Y")
    
    expediente = patient_data.get('Expediente', '')
    tipo_expediente = detectar_tipo_expediente(expediente)
    
    labs_faltantes_instruccion = f"""
INSTRUCCIÓN CRÍTICA SOBRE LABORATORIOS FALTANTES:
Tipo de expediente detectado: {tipo_expediente}

Si NO hay laboratorios de este mes:
- Expediente LE####: NO mencionar que no hay labs (omitir párrafo de labs en Análisis)
- Expediente RL####: Agregar "No se cuenta con estudios de laboratorio de este mes, se solicitan."
- Expediente BR####: Agregar "No se cuenta con estudios de laboratorio de este mes, se solicitan."
"""
    
    if labs_status == "OK":
        labs_section = f"""
═══════════════════════════════════════════════════════════════════════════════
DATOS DE LABORATORIO DEL CSV (USAR ESTOS VALORES)
═══════════════════════════════════════════════════════════════════════════════
{labs_data}

⚠️ INSTRUCCIÓN: Copia EXACTAMENTE la línea "Estudios de laboratorio:" de arriba en la nota.
NO agregues unidades (mg/dL, mmol/L, etc.) - el formato compacto ya está listo.
"""
    elif labs_status == "EXTRACT_FROM_HEMOHL7":
        labs_section = """
═══════════════════════════════════════════════════════════════════════════════
INSTRUCCIÓN SOBRE LABORATORIOS (MODO PDF CONSOLIDADO HEMOHL7)
═══════════════════════════════════════════════════════════════════════════════
Los laboratorios están en el reporte consolidado de HemoHL7 (ver sección abajo).

⚠️ EXTRAE TODOS los valores que encuentres y ponlos en sección "Estudios de laboratorio:":
Formato: Estudios de laboratorio: [Fecha]: Hb X g/dL, Hto X%, Cr X mg/dL, K X mEq/L, etc.
USA ABREVIATURAS. Esta sección es OBLIGATORIA.
"""
    elif labs_status == "EXTRACT_FROM_PDF_LABS":
        labs_section = """
═══════════════════════════════════════════════════════════════════════════════
INSTRUCCIÓN SOBRE LABORATORIOS (MODO PDFs DE LABS)
═══════════════════════════════════════════════════════════════════════════════
Los laboratorios están en PDFs de reportes de laboratorio (ver sección abajo).

⚠️ EXTRAE TODOS los valores que encuentres y ponlos en sección "Estudios de laboratorio:":
Formato: Estudios de laboratorio: [Fecha]: Hb X g/dL, Hto X%, Cr X mg/dL, K X mEq/L, etc.
USA ABREVIATURAS. Esta sección es OBLIGATORIA.
"""
    elif labs_status == "EXTRACT_FROM_NOTES":
        labs_section = """
═══════════════════════════════════════════════════════════════════════════════
INSTRUCCIÓN SOBRE LABORATORIOS (EXTRAER DE NOTAS HD)
═══════════════════════════════════════════════════════════════════════════════
NO hay CSV de laboratorios. BUSCA valores en las notas de hemodiálisis.

Si encuentras laboratorios en las notas HD:
- Formato: "Estudios de laboratorio: [Fecha]: Hb X, Hto X%, Cr X, K X, Ca X, P X, URR X%"
- SIN unidades (mg/dL, mmol/L, etc.)

Si NO hay laboratorios en las notas:
- CALCULAR URR estimado por parámetros de diálisis (ver tabla de cálculo)
- Escribir: "Estudios de laboratorio: No disponibles este mes. URR calculado por parámetros: ~XX%"
- En Análisis: "Adecuación dialítica estimada por parámetros con URR ~XX%"

⚠️ NUNCA escribir "se sugiere adecuación dialítica" - el Dr. ya lo hace constantemente.
"""
    else:
        labs_section = """
ESTUDIOS DE LABORATORIO: No disponibles este mes
"""
    
    if has_hd_notes:
        hd_section = f"""
NOTAS DE HEMODIÁLISIS DEL MES:
{hd_notes_text}
"""
    else:
        hd_section = """
NOTAS DE HEMODIÁLISIS: No disponibles en formato digital
"""
    
    if hemoHL7_text:
        hemoHL7_section = f"""

REPORTE CONSOLIDADO HEMOHL7:
{hemoHL7_text}
"""
    else:
        hemoHL7_section = ""
    
    if labs_pdf_text:
        labs_pdf_section = f"""

PDFs DE LABORATORIOS:
{labs_pdf_text}
"""
    else:
        labs_pdf_section = ""
    
    # v2.10.0: Construir sección de FN/Edad obligatoria si se conocen
    if fn_paciente and edad_paciente is not None:
        fn_edad_section = f"""

═══════════════════════════════════════════════════════════════════════════════
⚠️ FECHA DE NACIMIENTO Y EDAD - VALORES OBLIGATORIOS (NO INVENTAR)
═══════════════════════════════════════════════════════════════════════════════

Fecha de nacimiento: {fn_paciente}
Edad: {edad_paciente} años (calculada matemáticamente a la fecha de la nota)

⛔ INSTRUCCIONES CRÍTICAS:
- ESCRIBE EXACTAMENTE estos valores en el encabezado de la nota
- NO los recalcules, NO los modifiques, NO los infieras de otras fuentes
- Si las notas HD muestran una edad diferente, IGNORA esa edad
- En la sección Diagnósticos, usa "[Género] de {edad_paciente} años con los siguientes diagnósticos..."
- Estos valores ya fueron validados por el sistema
"""
    else:
        fn_edad_section = """

⚠️ FECHA DE NACIMIENTO: No registrada en caché ni encontrada en notas HD.
- Extrae la FN de las notas HD si está disponible
- Calcula la edad correctamente (considerar si ya pasó el cumpleaños)
- Si no hay FN disponible, omite ambos campos
"""
    
    prompt = f"""Eres un médico nefrólogo especialista. Genera una nota médica mensual de hemodiálisis siguiendo EXACTAMENTE la plantilla y reglas especificadas.

═══════════════════════════════════════════════════════════════════════════════
⛔ REGLA CRÍTICA #1 - NUNCA ESCRIBIR "NO DISPONIBLE"
═══════════════════════════════════════════════════════════════════════════════

ESTÁ ABSOLUTAMENTE PROHIBIDO escribir cualquiera de estas frases:
- "No disponible"
- "No especificado"  
- "No documentado"
- "No evaluado"
- "Sin información"
- "No se cuenta con datos"
- "Requiere evaluación"
- "Pendiente de..."

TODA la información que necesitas ESTÁ en las notas de hemodiálisis del mes (notas diarias).
Si no hay nota mensual previa, EXTRAE TODO de las notas diarias:
- Fecha de nacimiento: calcúlala de la edad que aparece en las notas
- Peso seco: aparece en TODAS las notas diarias
- Filtro: aparece en TODAS las notas diarias  
- Tiempo de sesión: aparece en TODAS las notas diarias
- QB, QD: aparecen en las notas diarias
- Acceso vascular: se menciona en las notas (catéter, fístula, etc.)
- Signos vitales: usa los del ÚLTIMO día del mes
- Medicamentos: se listan en las notas diarias
- EPO: dosis aparece en notas diarias

Si REALMENTE no encuentras un dato específico después de buscar en TODAS las notas:
- Para edad/fecha nacimiento: omite el campo
- Para otros datos: omite el dato o redacta de forma neutral sin inventar valores

═══════════════════════════════════════════════════════════════════════════════

DATOS DEL PACIENTE:
- Nombre: {patient_data['Paciente']}
- Expediente: {patient_data['Exped.2']}
- Procedencia: {patient_data['PROCED']}
- Sala: {patient_data['SALA']}
- Turno: {patient_data['TURNO']}
{fn_edad_section}
FECHA DE LA NOTA:
- Fecha de elaboración: {fecha_str}
- Mes de la nota: {mes_nota} {anio_nota}
- IMPORTANTE: Usa "{mes_nota}" en el título y referencias al mes

{labs_section}

{hemoHL7_section}

{labs_pdf_section}

{hd_section}

{labs_faltantes_instruccion}

═══════════════════════════════════════════════════════════════════════════════
PLANTILLA DE NOTA MÉDICA MENSUAL DE HEMODIÁLISIS
═══════════════════════════════════════════════════════════════════════════════

INSTRUCCIONES GENERALES:
- Rellena campos con información proporcionada del paciente
- No agregues ni infieras datos no explícitos
- MÁXIMO: 2 cuartillas y 4000 caracteres totales
- USAR ABREVIATURAS en estudios de laboratorio para ahorrar espacio
- Texto plano sin markdown (NO usar **, __, #, *)

═══════════════════════════════════════════════════════════════════════════════
REGLAS CRÍTICAS PARA DIAGNÓSTICOS (v2.8.0)
═══════════════════════════════════════════════════════════════════════════════

FECHAS DE DIAGNÓSTICOS:
- Si hay fecha exacta en las notas, ponerla entre paréntesis: "Diabetes mellitus tipo 2 (enero 2015)"
- Si dice "hace X años", CALCULAR el año: fecha actual {anio_nota} - X años = año de diagnóstico
  Ejemplo: "DM2 hace 8 años" → "Diabetes mellitus tipo 2 (año {anio_nota - 8})"
- Si no hay fecha disponible, NO inventar, omitir el paréntesis

ACCESO VASCULAR CON FECHA:
- Si hay fecha de instalación del acceso, incluirla: "Catéter tunelizado yugular derecho (18/02/2022)"
- Si hay fecha de creación de fístula: "FAVI radiocefálica izquierda (marzo 2020)"

HISTORIAL DE ACCESOS VASCULARES:
- Si las notas mencionan accesos vasculares previos, PRESERVAR esa información
- Ejemplo: "Antecedente de catéter temporal femoral derecho (2019), actualmente con FAVI radiocefálica izquierda (enero 2021)"
- NUNCA eliminar información de accesos previos, es clínicamente relevante

═══════════════════════════════════════════════════════════════════════════════
REGLAS CRÍTICAS PARA CANDIDATO A TRASPLANTE (v2.8.0)
═══════════════════════════════════════════════════════════════════════════════

Si el paciente NO es candidato a trasplante, DEBE incluirse la razón según guías KDIGO:

Formato: "Candidato a trasplante: No por [razón]."

RAZONES VÁLIDAS (basadas en KDIGO):
- Edad avanzada (>75 años con fragilidad)
- Comorbilidades cardiovasculares severas (ICC clase III-IV, cardiopatía isquémica activa)
- Neoplasia activa o reciente (<2 años libre de enfermedad)
- Obesidad mórbida (IMC >40 kg/m²)
- Enfermedad hepática avanzada
- Infección activa no controlada
- Demencia o deterioro cognitivo severo
- Mal apego al tratamiento documentado
- Enfermedad pulmonar severa (EPOC Gold IV)
- Diabetes con complicaciones severas múltiples
- Esperanza de vida limitada (<2 años)
- Condición psiquiátrica no controlada
- Abuso activo de sustancias

EJEMPLOS CORRECTOS:
- "Candidato a trasplante: No por edad avanzada y comorbilidades cardiovasculares."
- "Candidato a trasplante: No por obesidad mórbida y diabetes con complicaciones múltiples."
- "Candidato a trasplante: No por mal apego al tratamiento y condiciones socioeconómicas."
- "Candidato a trasplante: Sí, en protocolo de estudio."

SI ES CANDIDATO:
- "Candidato a trasplante: Sí." o "Candidato a trasplante: Sí, en protocolo."

═══════════════════════════════════════════════════════════════════════════════
ESTRUCTURA EXACTA DE LA NOTA
═══════════════════════════════════════════════════════════════════════════════

Encabezado demográfico (cada campo en línea separada):
Nombre: [Nombre completo en mayúsculas/título]
Fecha de nacimiento: [USAR el valor exacto de la sección "FECHA DE NACIMIENTO Y EDAD" arriba; si no se proporcionó, omitir]
Edad: [USAR el valor exacto de edad proporcionado; si no, calcular correctamente de FN considerando si ya cumplió años]
Inicio de hemodiálisis: [mes año]
Fecha de ingreso a Clínica: [dd/mm/aaaa]
Expediente: [número]
Fecha: {fecha_str}
Hora: [hora] hrs

Título centrado:
Nota Nefrología - Nota mensual {mes_nota}

FORMATO DE SECCIONES:
- TODO el contenido en la MISMA LÍNEA que el título (formato inline)
- INCLUYE Exploración física: debe ser "Exploración física: consciente, orientada..." en UNA sola línea
- SIN espacio/línea en blanco entre Signos vitales y Exploración física

Diagnósticos: [Género] de [edad] años con los siguientes diagnósticos: [INCLUIR TODOS los diagnósticos y antecedentes importantes de las notas previas]. [NO OMITIR: fracturas, paratiroidectomía, amputaciones, eventos cardiovasculares, cirugías, hospitalizaciones, enfermedad óseo-mineral, etc.]. [ERC G5 KDIGO en hemodiálisis desde mes/año]. Acceso vascular: [tipo, localización (fecha instalación si disponible)], [estado: funcional/disfuncional/permeable/semipermeable], [condiciones: limpio/signos de infección/parche seco]. [Historial de accesos previos si existe en notas]. Filtro: [modelo]. Sesiones por semana: [número]. Peso seco: [xx.x kg]. Ganancia interdialítica promedio: [rango mín-máx kg]. Eventos de bacteriemia: [Sí/No]. Alérgicos: [Negado/especificar]. Candidato a trasplante: [Sí/No por razón según KDIGO].

⚠️ ANTECEDENTES - MUY IMPORTANTE:
- COPIAR todos los diagnósticos que aparezcan en las notas previas del paciente
- Si la nota previa menciona "fractura", "paratiroidectomía", "amputación", "IAM", "EVC", etc. → INCLUIRLO
- NO simplificar ni resumir los antecedentes
- Si hay enfermedad óseo-mineral secundaria a ERC → INCLUIRLO

Evolución del mes: Durante {mes_nota} [adherencia a sesiones]. [SOLO eventos o síntomas POSITIVOS que ocurrieron, NO listar síntomas negativos]. [Estado del acceso vascular]. [Tolerancia a ultrafiltración]. [Eventos de hipotensión/hipertensión si ocurrieron].

⚠️ EVOLUCIÓN - NO ESCRIBIR:
- "acude asintomático"
- "niega náuseas, vómito, diarrea, fiebre, escalofríos..."
- Listas de síntomas que el paciente NO tiene
- Si no hubo eventos, solo decir "sin eventos adversos documentados"

Signos vitales: TA [/] mmHg, FC [x] lpm, FR [x] rpm, T [x] °C, SpO2 [x] %.
Exploración física: consciente, orientada, [exploración por sistemas relevante].

⚠️ SECCIÓN OBLIGATORIA - ESTUDIOS DE LABORATORIO:
Estudios de laboratorio: [Fecha]: Hb X, Hto X%, Leu X, PLQ X, Cr X, BUN X, Glu X, Na X, K X, Ca X, P X, Alb X, Ferritina X, URR X%.
- FORMATO COMPACTO: Solo valores numéricos, SIN unidades (mg/dL, mmol/L, etc.)
- Ejemplo correcto: "Hb 11.1, Hto 36.2%, Cr 7.38, K 5.53, Ca 8.5, P 4.2, Alb 3.8"
- USAR los valores EXACTOS del CSV si están disponibles
- Si NO hay laboratorios de este mes, ver instrucción de URR calculado abajo

Análisis: [Párrafo 1: Contexto clínico del paciente, adherencia, estado general]. [Párrafo 2: Interpretación de laboratorios relevantes, metas KDIGO]. [Estado del acceso vascular si aplica].
⚠️ Los valores de URR y otros labs en el Análisis DEBEN COINCIDIR con los de "Estudios de laboratorio"

Plan: [Continuidad del manejo. Ajustes de medicamentos. Vigilancia específica. Seguimiento de trasplante si aplica].

═══════════════════════════════════════════════════════════════════════════════
ABREVIATURAS OBLIGATORIAS (SIN UNIDADES para ahorrar espacio)
═══════════════════════════════════════════════════════════════════════════════

FORMATO COMPACTO - Solo abreviatura + valor numérico:
Hb 11.1, Hto 36.2%, Cr 7.38, K 5.53, Ca 8.5, P 4.2, CaxP 35.7, Na 138, Alb 3.8, Glu 95, BUN 45, PLQ 204, Leu 5.8, Ferritina 450, PTH 280, URR 72%

TABLA DE ABREVIATURAS:
Hemoglobina → Hb | Hematocrito → Hto | Creatinina → Cr | Albúmina → Alb
Sodio → Na | Potasio → K | Calcio → Ca | Fósforo → P | Cloro → Cl
Producto Calcio-Fósforo → CaxP (calculado automáticamente, NO modificar)
Glucosa → Glu | Leucocitos → Leu | Plaquetas → PLQ | Hierro → Fe
Colesterol → Col | Triglicéridos → TG | Ácido úrico → AU

═══════════════════════════════════════════════════════════════════════════════
URR CALCULADO (CUANDO NO HAY LABORATORIOS)
═══════════════════════════════════════════════════════════════════════════════

Si NO hay estudios de laboratorio de este mes pero hay parámetros de diálisis:
1. CALCULAR URR estimado basado en: filtro, QB (flujo sanguíneo), QD (flujo dializante), tiempo de sesión
2. Fórmula aproximada: URR aumenta con mayor QB, QD y tiempo
   - QB 300-350 + QD 500 + 3.5-4h + filtro alto flujo → URR ~65-75%
   - QB 250-300 + QD 500 + 3-3.5h + filtro bajo flujo → URR ~55-65%
3. ESCRIBIR: "URR calculado por parámetros de diálisis: ~XX%"
4. En Análisis mencionar: "Adecuación dialítica estimada por parámetros con URR ~XX%"

═══════════════════════════════════════════════════════════════════════════════
⛔ FRASES PROHIBIDAS - NO ESCRIBIR NUNCA
═══════════════════════════════════════════════════════════════════════════════

Esta nota es del Dr. Josué Tapia (nefrólogo). NO escribir frases que critiquen o sugieran que él no hace bien su trabajo:

❌ NUNCA ESCRIBIR:
- "Se sugiere adecuación dialítica" (él ya lo hace constantemente)
- "Se recomienda optimizar parámetros de diálisis" (él ya lo hace)
- "Considerar ajuste de tiempo/flujos" como sugerencia
- "Sería conveniente..." o "Se podría mejorar..."
- Cualquier frase que implique que el manejo actual es inadecuado

✅ EN CAMBIO ESCRIBIR:
- "Continuar parámetros de diálisis actuales" 
- "Mantener esquema de hemodiálisis"
- "Se mantiene adecuación dialítica con URR en meta"
- Si hay problema real documentado: describir el hallazgo y el PLAN de acción ya tomado

Tratamiento crónico: [Lista completa con DOSIS EXACTAS separadas por punto].

Parámetros de diálisis: [X] sesiones por semana, filtro [modelo], duración [X] min, QB [X] ml/min, QD [X] ml/min, peso seco [X] kg, temperatura [X] °C, heparina [bolo/infusión dosis], K [X], HCO3 [X], Ca [X], [sello con tipo si es catéter], EPO [tipo] [dosis] UI/semana.

Pronóstico: Funcional: Malo. Para la vida: Reservado. Candidato a trasplante: [Sí/No por razón].

═══════════════════════════════════════════════════════════════════════════════
FIRMA - IMPORTANTE: NO DUPLICAR
═══════════════════════════════════════════════════════════════════════════════

INSTRUCCIÓN CRÍTICA SOBRE LA FIRMA:
- La firma se agregará AUTOMÁTICAMENTE por el sistema
- NO incluyas la firma en tu respuesta
- NO escribas "Dr. Josué Wigberto Tapia López" al final
- NO escribas las cédulas profesionales
- Termina la nota en "Pronóstico:" y NADA MÁS

La firma será agregada automáticamente como:
Dr. Josué Wigberto Tapia López
Nefrólogo CP 9940966 SSA 5614 CMN 1267

═══════════════════════════════════════════════════════════════════════════════
VERIFICACIONES OBLIGATORIAS
═══════════════════════════════════════════════════════════════════════════════

⚠️ LÍMITE DE CARACTERES: MÁXIMO 4000 caracteres TOTALES ⚠️
- HemoHL7 corta a 4000 caracteres
- NO usar saltos de línea innecesarios
- USAR ABREVIATURAS en laboratorios
- Formato compacto: una línea por sección

CHECKLIST ANTES DE TERMINAR:
1. ✅ ¿Incluí sección "Estudios de laboratorio:" con valores del CSV? (OBLIGATORIO)
2. ✅ ¿Los valores de URR/Hb/Cr en Análisis coinciden con Estudios de laboratorio?
3. ✅ ¿Usé abreviaturas (Hb, Hto, Cr, K, P, Ca, etc.)?
4. ✅ ¿La nota tiene 4000 caracteres o menos?
5. ✅ ¿NO incluí firma al final?
6. ✅ ¿Texto plano sin markdown ni asteriscos?

═══════════════════════════════════════════════════════════════════════════════
CRITERIOS DE AUDITORÍA - CÉDULA DE EVALUACIÓN ALBA (v2.8.0)
═══════════════════════════════════════════════════════════════════════════════

La nota debe cumplir con los siguientes criterios de evaluación para auditoría:

**1. ADECUACIÓN DIALÍTICA:**
- Documentar Kt/V o URR si disponible (meta URR >65%, Kt/V >1.2)
- Mencionar tiempo efectivo de sesiones
- Correlación clínica si hay URR/Kt/V bajo

**2. ACCESO VASCULAR:**
- Tipo de acceso (FAVI preferente sobre catéter)
- Estado funcional (funcional/disfuncional/semipermeable)
- Plan de salida de catéter si aplica
- Detectar y documentar disfunción si existe

**3. ANEMIA Y HIERRO:**
- Hemoglobina con interpretación (meta 10-12 g/dL según KDIGO)
- Dosis de EPO actual
- Estado de hierro si disponible (ferritina, saturación)

**4. CKD-MBD (METABOLISMO ÓSEO-MINERAL):**
- Calcio, fósforo, PTH con interpretación
- Quelantes de fósforo en tratamiento si hiperfosfatemia
- Ajuste de calcitriol/análogos si aplica

**5. VOLUMEN Y HEMODINAMIA:**
- Peso seco documentado
- Ganancia interdialítica (meta <4% del peso seco)
- Eventos de hipotensión si ocurrieron
- Tolerancia a ultrafiltración

**6. SEGURIDAD DEL PACIENTE:**
- Eventos de bacteriemia documentados (Sí/No)
- Estado de infección del acceso vascular
- Alergias documentadas

**7. DOCUMENTACIÓN NORMATIVA (NOM-003-SSA3-2010):**
- Fecha y hora de la nota
- Diagnósticos completos con CIE-10 implícito
- Evolución del mes clara y concisa
- Signos vitales completos
- Exploración física pertinente
- Plan de manejo documentado
- Firma del médico (se agrega automáticamente)

**8. TRATAMIENTO DOCUMENTADO:**
- Medicamentos con dosis exactas
- Parámetros de diálisis completos (filtro, tiempo, flujos, baño)
- EPO con dosis semanal
- Sello de catéter si aplica

═══════════════════════════════════════════════════════════════════════════════
⛔ RANGOS DE REFERENCIA KDIGO PARA HEMODIÁLISIS - OBLIGATORIO
═══════════════════════════════════════════════════════════════════════════════

⚠️ USAR ESTOS RANGOS para interpretar laboratorios, NO rangos de población general:

ELECTROLITOS:
• Potasio (K): Meta 3.5-5.5 mEq/L
  - K 3.5-5.5 → "potasio en meta" (NO es hipopotasemia si K ≥3.5)
  - K <3.5 → hipopotasemia
  - K 5.5-6.5 → hiperpotasemia leve
  - K >6.5 → hiperpotasemia severa

• Sodio (Na): 135-145 mEq/L

• Calcio (Ca): Meta 8.4-9.5 mg/dL, tolerar hasta 10 mg/dL (KDIGO 2D)
  - Ca 8.4-10.0 → "calcio en rango aceptable" (NO es hipocalcemia si Ca ≥8.4)
  - Ca <8.4 → hipocalcemia
  - Ca >10.0 → hipercalcemia

• Fósforo (P): Meta 2.5-4.5 mg/dL, tolerar hasta 5 mg/dL en diálisis (KDIGO 2C)
  - P ≤5.0 → "fósforo en rango aceptable para diálisis" (NO es hiperfosfatemia si P ≤5.0)
  - P 5.0-6.5 → hiperfosfatemia leve
  - P >6.5 → hiperfosfatemia moderada-severa

METABOLISMO ÓSEO-MINERAL (CKD-MBD):
• PTH en hemodiálisis (G5D): Meta 150-300 pg/mL, evitar <100 o >500 (KDIGO 2B)
  - PTH <100 → PTH suprimida (riesgo enfermedad ósea adinámica)
  - PTH 100-500 → rango aceptable en diálisis
  - PTH >500 → hiperparatiroidismo secundario no controlado

• Calcidiol (25-OH vitamina D): Meta >30 ng/mL (KDIGO 2B)
  - <20 ng/mL → deficiencia
  - 20-30 ng/mL → insuficiencia  
  - >30 ng/mL → suficiente

• Producto Calcio-Fósforo (CaxP): Meta <55 mg²/dL² (KDIGO 2C)
  - CaxP <55 → "producto CaxP en meta, bajo riesgo de calcificación vascular"
  - CaxP 55-70 → "producto CaxP elevado, riesgo aumentado de calcificación vascular"
  - CaxP >70 → "producto CaxP críticamente elevado, alto riesgo de calcificación, ajustar quelantes y vitamina D"
  - El CaxP ya viene calculado en la línea de Estudios de laboratorio (ej: CaxP 35.7)
  - INCLUIRLO en el Análisis cuando esté disponible

ANEMIA Y HIERRO:
• Hemoglobina: Meta 10-12 g/dL
• Ferritina: Meta 200-500 ng/mL
  - <200 → depósitos bajos (indicar hierro IV)
  - 200-500 → adecuados
  - >800 → sobrecarga (suspender hierro)
• Saturación de transferrina: Meta 20-50%

NUTRICIÓN:
• Albúmina: Meta ≥3.5 g/dL
  - ≥3.5 → estado nutricional adecuado
  - 3.0-3.5 → hipoalbuminemia leve
  - <3.0 → desnutrición proteica

ADECUACIÓN DIALÍTICA:
• URR: Meta >65%
• Kt/V: Meta >1.2

⛔ REGLA CRÍTICA - NO DIAGNOSTICAR FALSAMENTE:
- K 3.8 → NO es hipopotasemia (está en meta 3.5-5.5)
- Ca 9.0 → NO es hipocalcemia (está en meta 8.4-10)
- P 4.8 → NO es hiperfosfatemia (tolerable hasta 5 en diálisis)
- PTH 280 → NO es hiperparatiroidismo (está en meta 150-300)
- CaxP 38 → NO es elevado (meta <55)

═══════════════════════════════════════════════════════════════════════════════
⚠️ LÓGICA CLÍNICA ESPECÍFICA - MUY IMPORTANTE
═══════════════════════════════════════════════════════════════════════════════

**INTERPRETACIÓN DE HEMOGLOBINA (metas KDIGO 10-12 g/dL):**

Si Hb > 12 g/dL:
- NO escribir "anemia no controlada" - NO tiene anemia
- Escribir: "Hemoglobina arriba de metas KDIGO"
- BUSCAR en notas diarias si se ajustó dosis de EPO (disminuyó o suspendió)
- Si hay ajuste documentado: "Se disminuyó/suspendió eritropoyetina por hemoglobina supraterapéutica"
- NO asumir "hemoconcentración" sin justificación clínica documentada (deshidratación, ultrafiltración excesiva)

Si Hb 10-12 g/dL:
- Escribir: "Hemoglobina en metas KDIGO (10-12 g/dL)"
- Sin ajuste de EPO necesario

Si Hb < 10 g/dL:
- Escribir: "Hemoglobina debajo de metas KDIGO"
- COMPARAR con estudios previos si los hay
- BUSCAR en notas si se incrementó EPO o se agregó hierro IV
- Si incrementó: "Se incrementó dosis de eritropoyetina" o "Se agregó hierro intravenoso"
- Si hay sangrado documentado (melena, sangrado GI, menstruación): mencionarlo como causa

**INTERPRETACIÓN DE AZOADOS (BUN, Creatinina):**

❌ NO escribir: "BUN y creatinina reflejan función renal residual mínima" (es obvio en ERC G5)

✅ Escribir según contexto:
- Si URR >65% o Kt/V >1.2: "Azoados en control con terapia sustitutiva"
- Si BUN >100 mg/dL: "Hiperazoemia severa, evaluar adecuación dialítica"
- Si Cr estable mes a mes: "Creatinina estable en contexto de terapia sustitutiva"

**SECCIÓN "PLAN" - NO REDUNDANCIA:**

❌ NO repetir en Plan lo que ya está en Parámetros de diálisis:
- NO: "Continuar hemodiálisis trisemanal 150 minutos, filtro 160, peso seco 36 kg"

✅ En Plan solo escribir:
- "Continuar terapia sustitutiva con parámetros actuales"
- Ajustes específicos de MEDICAMENTOS
- Vigilancias específicas (acceso vascular, infecciones, etc.)
- Seguimientos pendientes (valoraciones, trasplante, etc.)

**PARÁMETROS DE DIÁLISIS - INCLUIR PESO SECO:**

SIEMPRE incluir peso seco actual en la sección de Parámetros de diálisis:
"Parámetros de diálisis: 3 sesiones/semana, filtro FX120, 180 min, QB 350 ml/min, QD 500 ml/min, peso seco 36 kg, ..."

{NEFRONOTAS_MASTER_RULES}

GENERA AHORA la nota médica mensual con el lenguaje y reglas del GPT Nefro-Notas. IMPORTANTE: NO incluyas la firma al final."""

    return prompt


def postprocesar_nota_nefronotas(note_text):
    """Limpieza final de seguridad para alinear la salida con reglas Nefro-Notas."""
    if not note_text:
        return note_text

    # Eliminar markdown y encabezados accidentales
    note_text = note_text.replace('**', '').replace('__', '')
    note_text = re.sub(r'^#+\s*', '', note_text, flags=re.MULTILINE)


    # Sustituir nombre comercial frecuente por genérico
    note_text = re.sub(r'\bTums\b(?:\s*\d+\s*mg)?', 'carbonato de calcio 500 mg', note_text, flags=re.IGNORECASE)

    # Evitar sugerencias prohibidas comunes sobre K del dializado
    # Cubre verbos (disminuir/bajar/reducir), "K" o "potasio", y dializado/baño/líquido de diálisis
    note_text = re.sub(
        r'[^.\n]*(?:disminuir|bajar|reducir)[^.\n]*(?:\bK\b|potasio)[^.\n]*(?:dializad[oa]|baño|l[ií]quido de di[aá]lisis)[^.\n]*[.\n]?',
        '', note_text, flags=re.IGNORECASE)

    # Quitar firmas si el modelo las incluyó
    firma_patterns = [
        r'\n*Dr\.\s*Josué\s*Wigberto\s*Tapia\s*López.*?(?=\n|$)',
        r'\n*Nefrólogo\s*CP\s*\d+.*?(?=\n|$)',
        r'\n*CP\s*\d+\s*SSA\s*\d+\s*CMN\s*\d+.*?(?=\n|$)',
    ]
    for pattern in firma_patterns:
        note_text = re.sub(pattern, '', note_text, flags=re.IGNORECASE)

    # Compactar exceso de saltos sin unir secciones
    note_text = re.sub(r'\n{3,}', '\n\n', note_text)
    return note_text.strip()

def generate_note_with_claude(api_key, prompt, progress_callback=None):
    """Genera la nota usando la API de Claude"""
    try:
        client = anthropic.Anthropic(api_key=api_key)
        
        if progress_callback:
            progress_callback("Enviando solicitud a Claude API...")
        
        message = client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=8000,
            temperature=0.25,
            messages=[
                {"role": "user", "content": prompt}
            ]
        )
        
        if progress_callback:
            progress_callback("Nota generada exitosamente")
        
        note_text = message.content[0].text
        
        note_text = note_text.replace('**', '')
        note_text = note_text.replace('__', '')
        note_text = re.sub(r'^#+\s*', '', note_text, flags=re.MULTILINE)
        
        firma_patterns = [
            r'\n*Dr\.\s*Josué\s*Wigberto\s*Tapia\s*López.*?(?=\n|$)',
            r'\n*Nefrólogo\s*CP\s*\d+.*?(?=\n|$)',
            r'\n*CP\s*\d+\s*SSA\s*\d+\s*CMN\s*\d+.*?(?=\n|$)',
        ]
        
        for pattern in firma_patterns:
            note_text = re.sub(pattern, '', note_text, flags=re.IGNORECASE)
        
        note_text = postprocesar_nota_nefronotas(note_text)
        
        return note_text
    
    except anthropic.AuthenticationError:
        return "ERROR: API key inválida. Verifica tu clave de API."
    except anthropic.RateLimitError:
        return "ERROR: Límite de tasa excedido. Espera un momento e intenta nuevamente."
    except Exception as e:
        return f"ERROR: {str(e)}"

# ============================================================================
# v2.8.6: FUNCIÓN DE AJUSTE INTELIGENTE DE NOTAS LARGAS
# ============================================================================

def ajustar_nota_larga(api_key, nota_original, nombre_paciente, max_chars=3900):
    """
    Reajusta una nota que excede el límite de caracteres usando Claude.
    Mantiene TODA la información clínica esencial pero la hace más concisa.
    """
    try:
        client = anthropic.Anthropic(api_key=api_key)
        
        prompt_ajuste = f"""Eres un médico nefrólogo. La siguiente nota médica de "{nombre_paciente}" tiene {len(nota_original)} caracteres y excede el límite de 4000 caracteres para HemoHL7.

NOTA ORIGINAL:
{nota_original}

INSTRUCCIONES DE REAJUSTE:

1. Reescribe la nota para que tenga MÁXIMO {max_chars} caracteres

2. CONSERVA COMPLETO (NO reducir):
   - Todos los diagnósticos con fechas
   - TODOS los valores de laboratorio
   - Acceso vascular: tipo, estado, fecha
   - Medicamentos con dosis
   - Parámetros de diálisis
   - Pronóstico

3. ESTRATEGIAS para reducir:
   - Abreviaturas: Hb, Hto, Cr, K, P, Ca, Na, Leu, PLQ, Alb, Fe, PTH, Glu
   - Exploración física: solo hallazgos anormales
   - Evolución: eliminar frases redundantes
   - Análisis: condensar en 2-3 oraciones
   - Plan: frases cortas con punto

4. Mantener estructura de secciones
5. NO incluir firma
6. Texto plano sin markdown

Genera la nota ajustada:"""

        message = client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=4000,
            temperature=0.2,
            messages=[
                {"role": "user", "content": prompt_ajuste}
            ]
        )
        
        nota_ajustada = message.content[0].text
        
        nota_ajustada = nota_ajustada.replace('**', '')
        nota_ajustada = nota_ajustada.replace('__', '')
        nota_ajustada = re.sub(r'^#+\s*', '', nota_ajustada, flags=re.MULTILINE)
        
        firma_patterns = [
            r'\n*Dr\.\s*Josué\s*Wigberto\s*Tapia\s*López.*?(?=\n|$)',
            r'\n*Nefrólogo\s*CP\s*\d+.*?(?=\n|$)',
            r'\n*CP\s*\d+\s*SSA\s*\d+\s*CMN\s*\d+.*?(?=\n|$)',
        ]
        for pattern in firma_patterns:
            nota_ajustada = re.sub(pattern, '', nota_ajustada, flags=re.IGNORECASE)
        
        nota_ajustada = postprocesar_nota_nefronotas(nota_ajustada)
        
        return nota_ajustada, True
        
    except Exception as e:
        return nota_original[:max_chars], False

# ============================================================================

def create_word_document(patient_name, note_content):
    """Crea un documento Word con el formato EXACTO del Dr. Tapia - v2.8.1"""
    doc = Document()
    
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.57)
        section.bottom_margin = Inches(1.57)
        section.left_margin = Inches(0.79)
        section.right_margin = Inches(0.79)
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
    
    def to_title_case(text):
        """Convierte texto a formato título"""
        lowercase_words = {'de', 'del', 'la', 'las', 'los', 'y', 'e', 'o', 'u'}
        words = text.split()
        result = []
        for i, word in enumerate(words):
            if i == 0 or word.lower() not in lowercase_words:
                result.append(word.capitalize())
            else:
                result.append(word.lower())
        return ' '.join(result)
    
    lines = note_content.split('\n')
    
    sections_with_space_before = ['Diagnósticos', 'Evolución del mes', 'Signos vitales', 
                                   'Estudios de laboratorio', 
                                   'Análisis', 'Plan', 'Tratamiento crónico', 
                                   'Parámetros de diálisis', 'Pronóstico']
    
    last_section = None
    
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if not line:
            i += 1
            continue
        
        if line.startswith('Nombre:'):
            nombre_parte = line.split(':', 1)[1].strip()
            nombre_formateado = to_title_case(nombre_parte)
            line = f'Nombre: {nombre_formateado}'
        
        if any(line.startswith(x) for x in ['Nombre:', 'Fecha de nacimiento:', 'Edad:', 
                                              'Inicio de hemodiálisis:', 'Fecha de ingreso', 
                                              'Expediente:', 'Fecha:', 'Hora:']):
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.font.name = 'Century Gothic'
            run.font.size = Pt(10)
            run.bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.style = 'No Spacing'
            p.paragraph_format.line_spacing = 1.5
            i += 1
            continue
        
        if 'Nota Nefrología' in line or 'Nota mensual' in line:
            doc.add_paragraph()
            
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.font.name = 'Century Gothic'
            run.font.size = Pt(10)
            run.bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.style = 'No Spacing'
            p.paragraph_format.line_spacing = 1.5
            
            doc.add_paragraph()
            
            i += 1
            continue
        
        section_titles = [
            'Diagnósticos',
            'Evolución del mes',
            'Signos vitales',
            'Exploración física',
            'Estudios de laboratorio',
            'Análisis',
            'Plan',
            'Tratamiento crónico',
            'Parámetros de diálisis',
            'Pronóstico'
        ]
        
        is_section = False
        for section_title in section_titles:
            if line.startswith(section_title + ':'):
                is_section = True
                
                if section_title != 'Exploración física' and last_section is not None:
                    doc.add_paragraph()
                
                parts = line.split(':', 1)
                p = doc.add_paragraph()
                
                run_title = p.add_run(parts[0] + ':')
                run_title.font.name = 'Century Gothic'
                run_title.font.size = Pt(10)
                run_title.bold = True
                
                if len(parts) > 1 and parts[1].strip():
                    run_content = p.add_run('  ' + parts[1].strip())
                    run_content.font.name = 'Century Gothic'
                    run_content.font.size = Pt(10)
                    run_content.bold = False
                
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.style = 'No Spacing'
                p.paragraph_format.line_spacing = 1.5
                
                last_section = section_title
                break
        
        if is_section:
            i += 1
            continue
        
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.name = 'Century Gothic'
        run.font.size = Pt(10)
        run.bold = False
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.style = 'No Spacing'
        p.paragraph_format.line_spacing = 1.5
        
        i += 1
    
    doc.add_paragraph()
    
    p_firma = doc.add_paragraph()
    run_firma = p_firma.add_run('Dr. Josué Wigberto Tapia López')
    run_firma.font.name = 'Century Gothic'
    run_firma.font.size = Pt(10)
    run_firma.bold = True
    p_firma.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_firma.style = 'No Spacing'
    p_firma.paragraph_format.line_spacing = 1.5
    
    p_cedulas = doc.add_paragraph()
    run_cedulas = p_cedulas.add_run('Nefrólogo CP 9940966 SSA 5614 CMN 1267')
    run_cedulas.font.name = 'Century Gothic'
    run_cedulas.font.size = Pt(10)
    run_cedulas.bold = True
    p_cedulas.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cedulas.style = 'No Spacing'
    p_cedulas.paragraph_format.line_spacing = 1.5
    
    return doc

# ============================================================================
# v2.9.0: FUNCIONES DE ESTADO CON PERSISTENCIA A DISCO
# ============================================================================

def inicializar_estado_generacion():
    """Inicializa o recupera el estado de generación desde DISCO"""
    # Primero verificar si hay progreso en disco
    if hay_progreso_pendiente():
        progreso = cargar_progreso_de_disco()
        notas = cargar_notas_de_disco()
        
        if notas or progreso['completados']:
            st.session_state.partial_notes = notas
            st.session_state.processed_expedientes = progreso['completados']
            st.session_state.generation_errors = progreso['errores']
            st.session_state.generation_warnings = progreso['warnings']
            st.session_state.notas_ajustadas = progreso['ajustadas']
            st.session_state.progreso_recuperado_disco = True
    
    # Inicializar valores por defecto si no existen
    if 'partial_notes' not in st.session_state:
        st.session_state.partial_notes = []
    if 'processed_expedientes' not in st.session_state:
        st.session_state.processed_expedientes = set()
    if 'generation_errors' not in st.session_state:
        st.session_state.generation_errors = []
    if 'generation_warnings' not in st.session_state:
        st.session_state.generation_warnings = []
    if 'notas_ajustadas' not in st.session_state:
        st.session_state.notas_ajustadas = []
    if 'generation_in_progress' not in st.session_state:
        st.session_state.generation_in_progress = False
    if 'last_generated_time' not in st.session_state:
        st.session_state.last_generated_time = None
    if 'progreso_recuperado_disco' not in st.session_state:
        st.session_state.progreso_recuperado_disco = False

def crear_zip_parcial(notas):
    """Crea un ZIP con las notas proporcionadas"""
    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
        for note in notas:
            filename = f"{note['expediente']}_{note['paciente'].replace(' ', '_')}.docx"
            zip_file.writestr(filename, note['doc_bytes'])
    zip_buffer.seek(0)
    return zip_buffer

def limpiar_estado_generacion():
    """Limpia el estado de generación para empezar de nuevo - INCLUYE DISCO"""
    st.session_state.partial_notes = []
    st.session_state.processed_expedientes = set()
    st.session_state.generation_errors = []
    st.session_state.generation_warnings = []
    st.session_state.notas_ajustadas = []
    st.session_state.generation_in_progress = False
    st.session_state.progreso_recuperado_disco = False
    
    # Limpiar también de disco
    limpiar_progreso_de_disco()

# ============================================================================
# INTERFAZ PRINCIPAL
# ============================================================================

def main():
    st.markdown('<h1 class="main-header">🩺 NefroNotes Batch Generator v2.12.0</h1>', unsafe_allow_html=True)
    st.markdown("""
    <div style='text-align: center; color: #6B7280; margin-bottom: 2rem;'>
        Sistema de generación masiva de notas nefrológicas mensuales<br>
        <b>Dr. Josué Wigberto Tapia López</b> - Centro Médico Nacional del Bajío, IMSS
    </div>
    """, unsafe_allow_html=True)
    
    # Inicializar estado (ahora carga desde disco si hay progreso)
    inicializar_estado_generacion()
    
    with st.sidebar:
        st.image("https://img.icons8.com/fluency/96/medical-heart.png", width=80)
        st.title("⚙️ Configuración")
        
        api_key = st.text_input(
            "🔑 API Key de Claude",
            type="password",
            help="Obtén tu API key en console.anthropic.com"
        )
        
        if not api_key:
            st.warning("⚠️ Necesitas una API key para continuar")
            st.markdown("""
            **Cómo obtener tu API key:**
            1. Ve a [console.anthropic.com](https://console.anthropic.com)
            2. Crea una cuenta o inicia sesión
            3. Compra $5 USD en créditos
            4. Ve a "API Keys" y crea una nueva
            """)
        
        st.divider()
        
        st.markdown("### 📅 Fecha de la Nota")
        usar_fecha_personalizada = st.checkbox(
            "Usar fecha personalizada",
            help="Marca esta opción para generar notas con una fecha específica (útil para cortes antes de fin de mes)"
        )
        
        if usar_fecha_personalizada:
            fecha_nota = st.date_input(
                "Fecha de elaboración",
                value=date.today(),
                help="Selecciona la fecha que aparecerá en las notas"
            )
            st.info(f"📆 Las notas se generarán con fecha: **{fecha_nota.strftime('%d/%m/%Y')}** (mes de {MESES_ESP[fecha_nota.month]})")
        else:
            fecha_nota = None
            st.caption("Se usará la fecha actual automáticamente")
        
        st.divider()
        
        # v2.9.0: Estadísticas mejoradas con indicador de disco
        st.markdown("### 📊 Estadísticas de Sesión")
        col_stat1, col_stat2 = st.columns(2)
        with col_stat1:
            st.metric("✅ Generadas", len(st.session_state.partial_notes))
        with col_stat2:
            st.metric("❌ Errores", len(st.session_state.generation_errors))
        
        # v2.9.0: Mostrar si hay progreso guardado en disco
        if st.session_state.partial_notes:
            if st.session_state.get('progreso_recuperado_disco'):
                st.success(f"💾 {len(st.session_state.partial_notes)} notas (recuperado de disco)")
            else:
                st.success(f"💾 {len(st.session_state.partial_notes)} notas guardadas")
        
        # v2.9.0: Estado del disco
        if hay_progreso_pendiente():
            resumen = obtener_resumen_progreso()
            if resumen:
                st.info(f"💿 Disco: {resumen['notas']} notas guardadas")
    
    st.session_state.fecha_nota = fecha_nota
    
    tab1, tab2, tab3, tab4 = st.tabs(["📤 Cargar Datos", "⚡ Generar Notas", "📥 Descargar Resultados", "📅 Fechas de Nacimiento"])
    
    # ========================================================================
    # TAB 1: CARGAR DATOS
    # ========================================================================
    with tab1:
        st.markdown('<div class="step-header"><h2>Paso 1: Carga tus archivos</h2></div>', unsafe_allow_html=True)
        
        col1, col2, col3 = st.columns(3)
        
        with col1:
            st.subheader("📋 Censo de Pacientes")
            censo_file = st.file_uploader(
                "Archivo Excel (.xlsx)",
                type=['xlsx'],
                key='censo',
                help="MÍNIMO 2 columnas: Expediente y Nombre"
            )
            
            if censo_file:
                try:
                    censo_df_raw = pd.read_excel(censo_file)
                    censo_df = normalize_censo_columns(censo_df_raw)
                    st.success(f"✅ {len(censo_df)} pacientes cargados")
                    st.session_state.censo_df = censo_df
                    
                    with st.expander("Ver muestra del censo"):
                        st.dataframe(censo_df.head())
                except Exception as e:
                    st.error(f"Error al leer censo: {str(e)}")
        
        with col2:
            st.subheader("🧪 Laboratorios (Opcional)")
            labs_file = st.file_uploader(
                "Archivo CSV (.csv) - OPCIONAL",
                type=['csv'],
                key='labs',
                help="Solo para Alba Centro (LE)"
            )
            
            if labs_file:
                try:
                    labs_df = pd.read_csv(labs_file, encoding='latin1')
                    st.success(f"✅ {len(labs_df)} registros")
                    st.session_state.labs_df = labs_df
                    
                    # v2.8.9: Diagnóstico de matching
                    if 'censo_df' in st.session_state:
                        with st.expander("🔍 Diagnóstico de matching CSV-Censo"):
                            censo = st.session_state.censo_df
                            matches_con_labs = 0
                            matches_sin_labs = 0
                            no_matches = []
                            
                            import math
                            
                            # Detectar columna de nombre en censo
                            nombre_col_censo = None
                            for col in censo.columns:
                                if 'paciente' in str(col).lower():
                                    nombre_col_censo = col
                                    break
                            if nombre_col_censo is None:
                                nombre_col_censo = censo.columns[1] if len(censo.columns) > 1 else censo.columns[0]
                            
                            # Crear set de nombres normalizados del CSV
                            labs_df['_temp_norm'] = labs_df['PatName'].apply(normalizar_nombre)
                            
                            for _, row in censo.iterrows():
                                nombre_censo = str(row[nombre_col_censo])
                                nombre_norm = normalizar_nombre(nombre_censo)
                                
                                # Buscar coincidencia
                                found = labs_df[labs_df['_temp_norm'] == nombre_norm]
                                
                                if found.empty:
                                    apellidos = tuple(nombre_norm.split()[:2])
                                    for idx, lab_row in labs_df.iterrows():
                                        if apellidos == tuple(lab_row['_temp_norm'].split()[:2]):
                                            found = labs_df.iloc[[idx]]
                                            break
                                
                                if not found.empty:
                                    # Verificar si tiene valores válidos
                                    hgb = found.iloc[0].get('Hgb')
                                    if hgb is not None and not (isinstance(hgb, float) and math.isnan(hgb)):
                                        matches_con_labs += 1
                                    else:
                                        matches_sin_labs += 1
                                else:
                                    no_matches.append(nombre_censo)
                            
                            labs_df.drop('_temp_norm', axis=1, inplace=True)
                            
                            st.write(f"✅ **{matches_con_labs}** pacientes CON laboratorios válidos en CSV")
                            st.write(f"⚠️ **{matches_sin_labs}** pacientes en CSV pero SIN valores (se extraerá de notas HD)")
                            st.write(f"❌ **{len(no_matches)}** pacientes NO encontrados en CSV")
                            
                            if no_matches and len(no_matches) <= 5:
                                st.caption("No encontrados: " + ", ".join(no_matches[:5]))
                except Exception as e:
                    st.error(f"Error: {str(e)}")
                    import traceback
                    st.code(traceback.format_exc())
            else:
                st.info("ℹ️ Los labs se extraerán de PDFs o notas HD")
        
        with col3:
            st.subheader("📄 Notas de HD")
            hd_notes_files = st.file_uploader(
                "PDFs de notas HD (múltiples)",
                type=['pdf'],
                accept_multiple_files=True,
                key='hd_notes'
            )
            
            if hd_notes_files:
                st.success(f"✅ {len(hd_notes_files)} PDFs cargados")
                st.session_state.hd_notes_files = hd_notes_files
                
                hd_notes_dict = {}
                pdfs_sin_expediente = []
                
                for pdf_file in hd_notes_files:
                    full_path = pdf_file.name
                    expediente_match = re.search(r'[A-Z]{2}\d+', full_path, re.IGNORECASE)
                    
                    if expediente_match:
                        expediente = expediente_match.group().upper()
                        if expediente not in hd_notes_dict:
                            hd_notes_dict[expediente] = []
                        
                        text = extract_text_from_pdf(pdf_file)
                        hd_notes_dict[expediente].append({
                            'filename': full_path.split('/')[-1] if '/' in full_path else full_path,
                            'text': text
                        })
                    else:
                        pdfs_sin_expediente.append(full_path)
                
                st.session_state.hd_notes_dict = hd_notes_dict
                
                with st.expander(f"✅ PDFs organizados ({len(hd_notes_dict)} pacientes)"):
                    for exp, notes in sorted(hd_notes_dict.items()):
                        st.write(f"**{exp}**: {len(notes)} notas")
        
        st.markdown("---")
        st.subheader("📊 PDF Consolidado de Laboratorios (Opcional)")
        
        hemoHL7_consolidated_file = st.file_uploader(
            "PDF consolidado de laboratorios",
            type=['pdf'],
            key='hemoHL7_consolidated'
        )
        
        if hemoHL7_consolidated_file:
            with st.spinner("Procesando PDF consolidado..."):
                hemoHL7_labs_dict = parse_hemoHL7_consolidated_pdf(hemoHL7_consolidated_file)
                if hemoHL7_labs_dict:
                    st.success(f"✅ Laboratorios de {len(hemoHL7_labs_dict)} pacientes")
                    st.session_state.hemoHL7_labs_dict = hemoHL7_labs_dict
        
        st.markdown("---")
        st.subheader("🧪 PDFs de Laboratorios Individuales (Opcional)")
        
        labs_pdf_files = st.file_uploader(
            "PDFs de laboratorios (múltiples)",
            type=['pdf'],
            accept_multiple_files=True,
            key='labs_pdfs'
        )
        
        if labs_pdf_files:
            st.success(f"✅ {len(labs_pdf_files)} PDFs de labs")
            st.session_state.labs_pdf_files = labs_pdf_files
            
            labs_pdf_dict = {}
            for pdf_file in labs_pdf_files:
                full_path = pdf_file.name
                expediente_match = re.search(r'[A-Z]{2}\d+', full_path, re.IGNORECASE)
                
                if expediente_match:
                    expediente = expediente_match.group().upper()
                    if expediente not in labs_pdf_dict:
                        labs_pdf_dict[expediente] = []
                    
                    text = extract_text_from_pdf(pdf_file)
                    labs_pdf_dict[expediente].append({
                        'filename': full_path.split('/')[-1] if '/' in full_path else full_path,
                        'text': text
                    })
            
            st.session_state.labs_pdf_dict = labs_pdf_dict

    # ========================================================================
    # TAB 2: GENERAR NOTAS
    # ========================================================================
    with tab2:
        st.markdown('<div class="step-header"><h2>Paso 2: Generar Notas Médicas</h2></div>', unsafe_allow_html=True)
        
        prerequisites_met = (
            'censo_df' in st.session_state and
            api_key
        )
        
        if not prerequisites_met:
            st.warning("⚠️ Completa el Paso 1 y configura tu API key")
            return
        
        censo_df = st.session_state.censo_df
        labs_df = st.session_state.get('labs_df', None)
        hd_notes_dict = st.session_state.get('hd_notes_dict', {})
        labs_pdf_dict = st.session_state.get('labs_pdf_dict', {})
        hemoHL7_labs_dict = st.session_state.get('hemoHL7_labs_dict', {})
        fecha_nota = st.session_state.get('fecha_nota', None)
        
        if fecha_nota:
            st.info(f"📅 **Fecha de las notas:** {fecha_nota.strftime('%d/%m/%Y')} (mes de {MESES_ESP[fecha_nota.month]})")
        else:
            st.info(f"📅 **Fecha de las notas:** {datetime.now().strftime('%d/%m/%Y')} (fecha actual)")
        
        # ================================================================
        # v2.9.0: MOSTRAR PROGRESO RECUPERADO DE DISCO
        # ================================================================
        if st.session_state.partial_notes:
            total_censo = len(censo_df)
            generadas = len(st.session_state.partial_notes)
            faltantes = total_censo - generadas
            
            # Determinar si viene de disco
            if st.session_state.get('progreso_recuperado_disco'):
                st.markdown(f"""
                <div class="disk-saved-box">
                    <h4>💾 Progreso Recuperado de Disco</h4>
                    <p>✅ <b>{generadas}</b> notas generadas de <b>{total_censo}</b> ({faltantes} pendientes)</p>
                    <p><small>Si se cortó la conexión, tus notas están seguras en disco.</small></p>
                </div>
                """, unsafe_allow_html=True)
            else:
                st.markdown(f"""
                <div class="progress-box">
                    <h4>💾 Progreso Guardado</h4>
                    <p>✅ <b>{generadas}</b> notas generadas de <b>{total_censo}</b> ({faltantes} pendientes)</p>
                </div>
                """, unsafe_allow_html=True)
            
            col_btn1, col_btn2, col_btn3 = st.columns(3)
            
            with col_btn1:
                if st.button("▶️ Continuar donde se quedó", type="primary", use_container_width=True):
                    st.session_state.continue_generation = True
                    st.rerun()
            
            with col_btn2:
                # Botón de descarga parcial
                zip_parcial = crear_zip_parcial(st.session_state.partial_notes)
                st.download_button(
                    label=f"📦 Descargar {generadas} notas",
                    data=zip_parcial,
                    file_name=f"notas_parcial_{generadas}_de_{total_censo}.zip",
                    mime="application/zip",
                    use_container_width=True
                )
            
            with col_btn3:
                if st.button("🗑️ Borrar todo y reiniciar", use_container_width=True):
                    limpiar_estado_generacion()
                    st.rerun()
            
            st.markdown("---")
        
        # Opciones de generación
        col1, col2 = st.columns(2)
        
        with col1:
            generation_mode = st.radio(
                "Modo de generación",
                ["Generar todas las notas", "Seleccionar pacientes específicos"]
            )
        
        with col2:
            delay_between_requests = st.slider(
                "Delay entre solicitudes (segundos)",
                min_value=1,
                max_value=10,
                value=2,
                help="Recomendado: 2 segundos"
            )
        
        # Determinar pacientes a procesar
        if generation_mode == "Seleccionar pacientes específicos":
            patient_options = []
            for _, row in censo_df.iterrows():
                exp = row.get('Exped.2', row.get('Expediente', ''))
                nombre = row.get('Paciente', '')
                # Marcar los ya procesados
                marca = "✅ " if exp in st.session_state.processed_expedientes else ""
                option = f"{marca}{exp} - {nombre}"
                patient_options.append(option)
            
            selected_options = st.multiselect(
                "Selecciona los pacientes",
                options=patient_options,
                help="✅ = ya generada. Busca por expediente o nombre"
            )
            
            # Extraer los expedientes seleccionados (quitando la marca ✅)
            selected_expedientes = [opt.replace("✅ ", "").split(' - ')[0] for opt in selected_options]
            patients_to_process = censo_df[censo_df['Exped.2'].isin(selected_expedientes)]
        else:
            patients_to_process = censo_df
        
        # Calcular cuántos faltan (excluyendo ya procesados)
        expedientes_pendientes = [
            exp for exp in patients_to_process['Exped.2'].tolist() 
            if exp not in st.session_state.processed_expedientes
        ]
        
        total_a_generar = len(expedientes_pendientes)
        ya_generadas = len(st.session_state.processed_expedientes)
        
        if ya_generadas > 0:
            st.info(f"📊 Se generarán **{total_a_generar}** notas nuevas ({ya_generadas} ya completadas)")
        else:
            st.info(f"📊 Se generarán **{total_a_generar}** notas")
        
        # v2.8.9: VERIFICACIÓN PREVIA DE NOTAS HD
        pacientes_sin_notas = []
        pacientes_con_notas = []
        for exp in expedientes_pendientes:
            if exp in hd_notes_dict:
                pacientes_con_notas.append(exp)
            else:
                pacientes_sin_notas.append(exp)
        
        if pacientes_sin_notas:
            with st.expander(f"⚠️ {len(pacientes_sin_notas)} pacientes SIN notas HD (serán omitidos)", expanded=True):
                st.warning(f"""
                **Estos pacientes NO tienen PDFs de notas de hemodiálisis cargados:**
                
                {', '.join(pacientes_sin_notas[:20])}{'...' if len(pacientes_sin_notas) > 20 else ''}
                
                **Se omitirán** porque no se puede generar una nota sin información clínica.
                
                👉 **Solución:** Carga los PDFs faltantes en el Paso 1 y vuelve a intentar.
                """)
            
            st.info(f"✅ Se generarán **{len(pacientes_con_notas)}** notas (pacientes CON notas HD)")
        
        # ================================================================
        # BOTÓN DE GENERACIÓN
        # ================================================================
        should_generate = st.button("🚀 Iniciar Generación de Notas", type="primary", use_container_width=True)
        
        # También iniciar si se presionó "Continuar"
        if st.session_state.get('continue_generation', False):
            should_generate = True
            st.session_state.continue_generation = False
        
        if should_generate and total_a_generar > 0:
            
            progress_container = st.container()
            
            with progress_container:
                progress_bar = st.progress(0)
                status_text = st.empty()
                time_estimate = st.empty()
                results_container = st.container()
                
                # Placeholder para descarga parcial durante generación
                partial_download_placeholder = st.empty()
            
            # Filtrar solo los pacientes pendientes
            patients_pending = patients_to_process[
                ~patients_to_process['Exped.2'].isin(st.session_state.processed_expedientes)
            ]
            
            total_patients = len(patients_pending)
            start_time = time.time()
            
            for idx, (_, patient) in enumerate(patients_pending.iterrows()):
                try:
                    # Actualizar progreso
                    progress = (idx + 1) / total_patients
                    progress_bar.progress(progress)
                    
                    # Calcular tiempo estimado
                    elapsed = time.time() - start_time
                    if idx > 0:
                        avg_time = elapsed / idx
                        remaining = avg_time * (total_patients - idx)
                        mins_remaining = int(remaining // 60)
                        secs_remaining = int(remaining % 60)
                        time_estimate.write(f"⏱️ Tiempo estimado restante: {mins_remaining}m {secs_remaining}s")
                    
                    status_text.write(f"Procesando {idx + 1}/{total_patients}: **{patient['Paciente']}** ({patient['Exped.2']})")
                    
                    expediente = patient['Exped.2']
                    
                    # Notas HD
                    hd_notes_text = ""
                    has_hd_notes = False
                    
                    if expediente in hd_notes_dict:
                        has_hd_notes = True
                        for note in hd_notes_dict[expediente]:
                            hd_notes_text += f"\n--- {note['filename']} ---\n"
                            hd_notes_text += note['text']
                    
                    # v2.8.9: VALIDACIÓN CRÍTICA - No generar sin notas HD
                    if not has_hd_notes:
                        error_info = {
                            'paciente': patient['Paciente'],
                            'expediente': expediente,
                            'error': 'SIN NOTAS HD - No se encontraron PDFs de notas de hemodiálisis para este paciente'
                        }
                        st.session_state.generation_errors.append(error_info)
                        # v2.9.0: Guardar error a disco también
                        guardar_progreso_a_disco(None, errores=st.session_state.generation_errors)
                        progress_bar.progress((idx + 1) / total_patients)
                        continue  # Saltar este paciente
                    
                    # PDFs de labs
                    labs_pdf_text = ""
                    has_labs_pdf = False
                    
                    if expediente in labs_pdf_dict:
                        has_labs_pdf = True
                        for pdf in labs_pdf_dict[expediente]:
                            labs_pdf_text += f"\n--- {pdf['filename']} ---\n"
                            labs_pdf_text += pdf['text']
                    
                    # HemoHL7
                    hemoHL7_text = ""
                    has_hemoHL7 = False
                    
                    if expediente in hemoHL7_labs_dict:
                        has_hemoHL7 = True
                        hemoHL7_text = hemoHL7_labs_dict[expediente]
                    
                    # Procesar labs - v2.8.8: incluye nombre para match por PatName
                    labs_data, labs_status = process_labs_data(
                        labs_df, 
                        expediente,
                        patient['Paciente'],  # Nombre para match alternativo
                        hd_notes_text if has_hd_notes else None,
                        labs_pdf_text if has_labs_pdf else None,
                        hemoHL7_text if has_hemoHL7 else None
                    )
                    
                    # v2.10.0: Obtener FN y edad calculada
                    fecha_ref_edad = fecha_nota if fecha_nota else date.today()
                    fn_paciente, edad_paciente, origen_fn = obtener_fn_y_edad(
                        expediente, 
                        hd_notes_text if has_hd_notes else "", 
                        fecha_ref_edad
                    )
                    
                    # Si no se encontró, registrar warning para que el usuario lo agregue manualmente
                    if not fn_paciente:
                        st.session_state.generation_warnings.append({
                            'paciente': patient['Paciente'],
                            'expediente': expediente,
                            'mensaje': "⚠️ Sin FN: no se encontró en caché ni en notas HD. Agrégala manualmente en la pestaña 'Fechas de Nacimiento'."
                        })
                    
                    # Crear prompt
                    prompt = create_master_prompt(
                        patient, 
                        labs_data if labs_data else "", 
                        labs_status,
                        hd_notes_text if has_hd_notes else "",
                        has_hd_notes,
                        labs_pdf_text if has_labs_pdf else "",
                        hemoHL7_text if has_hemoHL7 else "",
                        fecha_nota,
                        fn_paciente,        # v2.10.0
                        edad_paciente       # v2.10.0
                    )
                    
                    # Generar nota
                    note_content = generate_note_with_claude(api_key, prompt)
                    
                    if note_content.startswith("ERROR"):
                        error_info = {
                            'paciente': patient['Paciente'],
                            'expediente': expediente,
                            'error': note_content
                        }
                        st.session_state.generation_errors.append(error_info)
                        # v2.9.0: Guardar error a disco
                        guardar_progreso_a_disco(None, errores=st.session_state.generation_errors)
                    else:
                        # Ajuste inteligente si es muy larga
                        note_length = len(note_content)
                        fue_ajustada = False
                        
                        if note_length > 4000:
                            status_text.write(f"⚙️ Ajustando nota de {patient['Paciente']} ({note_length} chars)...")
                            
                            nota_ajustada, exito_ajuste = ajustar_nota_larga(
                                api_key, 
                                note_content, 
                                patient['Paciente'],
                                max_chars=3900
                            )
                            
                            if exito_ajuste:
                                note_content = nota_ajustada
                                fue_ajustada = True
                                nuevo_length = len(note_content)
                                
                                st.session_state.notas_ajustadas.append({
                                    'paciente': patient['Paciente'],
                                    'expediente': expediente,
                                    'original': note_length,
                                    'ajustada': nuevo_length
                                })
                                
                                if nuevo_length > 4000:
                                    st.session_state.generation_warnings.append({
                                        'paciente': patient['Paciente'],
                                        'expediente': expediente,
                                        'mensaje': f"⚠️ Nota aún larga después de ajuste ({nuevo_length} chars)"
                                    })
                            else:
                                st.session_state.generation_warnings.append({
                                    'paciente': patient['Paciente'],
                                    'expediente': expediente,
                                    'mensaje': f"⚠️ Ajuste falló, truncada de {note_length} a 3900 chars"
                                })
                        
                        elif note_length > 3900:
                            st.session_state.generation_warnings.append({
                                'paciente': patient['Paciente'],
                                'expediente': expediente,
                                'mensaje': f"⚠️ Nota cerca del límite ({note_length} caracteres)"
                            })
                        
                        # v2.10.0: Validar y corregir FN/Edad post-generación
                        if fn_paciente and edad_paciente is not None:
                            note_content, fue_corregida_fn = validar_y_corregir_fn_edad_en_nota(
                                note_content, fn_paciente, edad_paciente
                            )
                            if fue_corregida_fn:
                                st.session_state.generation_warnings.append({
                                    'paciente': patient['Paciente'],
                                    'expediente': expediente,
                                    'mensaje': f"✏️ FN/Edad corregida automáticamente (FN: {fn_paciente}, Edad: {edad_paciente})"
                                })
                        
                        # Crear documento Word
                        doc = create_word_document(patient['Paciente'], note_content)
                        
                        doc_io = io.BytesIO()
                        doc.save(doc_io)
                        doc_io.seek(0)
                        
                        # Guardar nota
                        nota_generada = {
                            'paciente': patient['Paciente'],
                            'expediente': expediente,
                            'content': note_content,
                            'doc_bytes': doc_io.getvalue(),
                            'tiene_labs': labs_status == "OK",
                            'tiene_hd_notes': has_hd_notes,
                            'fue_ajustada': fue_ajustada
                        }
                        
                        st.session_state.partial_notes.append(nota_generada)
                        st.session_state.processed_expedientes.add(expediente)
                        
                        # ============================================================
                        # v2.9.0: GUARDAR A DISCO INMEDIATAMENTE DESPUÉS DE CADA NOTA
                        # ============================================================
                        guardar_nota_a_disco(nota_generada)
                        guardar_progreso_a_disco(
                            expediente,
                            errores=st.session_state.generation_errors,
                            warnings=st.session_state.generation_warnings,
                            ajustadas=st.session_state.notas_ajustadas
                        )
                    
                    # ============================================================
                    # v2.9.0: MOSTRAR GUARDADO CADA 5 NOTAS
                    # ============================================================
                    if (idx + 1) % 5 == 0:
                        notas_guardadas = len(st.session_state.partial_notes)
                        with partial_download_placeholder:
                            st.success(f"💾 Guardado en disco: {notas_guardadas} notas. Si se corta, el progreso está seguro.")
                    
                    # Delay entre requests
                    if idx < total_patients - 1:
                        time.sleep(delay_between_requests)
                
                except Exception as e:
                    error_info = {
                        'paciente': patient['Paciente'],
                        'expediente': patient['Exped.2'],
                        'error': str(e)
                    }
                    st.session_state.generation_errors.append(error_info)
                    guardar_progreso_a_disco(None, errores=st.session_state.generation_errors)
            
            # Finalización
            st.session_state.final_notes = st.session_state.partial_notes
            st.session_state.last_generated_time = datetime.now()
            
            with results_container:
                total_generadas = len(st.session_state.partial_notes)
                total_errores = len(st.session_state.generation_errors)
                
                st.success(f"✅ **Generación completada:** {total_generadas} notas, {total_errores} errores")
                
                # Tiempo total
                total_time = time.time() - start_time
                mins = int(total_time // 60)
                secs = int(total_time % 60)
                st.info(f"⏱️ Tiempo total: {mins}m {secs}s")
                
                # v2.9.0: Nota sobre persistencia
                st.success("💾 Todas las notas están guardadas en disco y listas para descargar.")
                
                if st.session_state.notas_ajustadas:
                    with st.expander(f"⚙️ Notas ajustadas automáticamente ({len(st.session_state.notas_ajustadas)})"):
                        for aj in st.session_state.notas_ajustadas:
                            st.info(f"**{aj['paciente']}** ({aj['expediente']}): {aj['original']} → {aj['ajustada']} chars")
                
                if st.session_state.generation_warnings:
                    with st.expander(f"⚠️ Ver advertencias ({len(st.session_state.generation_warnings)})"):
                        for warn in st.session_state.generation_warnings:
                            st.warning(f"**{warn['paciente']}** ({warn['expediente']}): {warn['mensaje']}")
                
                if st.session_state.generation_errors:
                    with st.expander("❌ Ver errores"):
                        for error in st.session_state.generation_errors:
                            st.error(f"**{error['paciente']}**: {error['error']}")
        
        elif should_generate and total_a_generar == 0:
            st.success("✅ ¡Todas las notas ya están generadas!")

    # ========================================================================
    # TAB 3: DESCARGAR RESULTADOS
    # ========================================================================
    with tab3:
        st.markdown('<div class="step-header"><h2>Paso 3: Descargar Notas Generadas</h2></div>', unsafe_allow_html=True)
        
        # v2.9.0: Cargar de disco si no hay en memoria
        available_notes = st.session_state.partial_notes
        if not available_notes and hay_progreso_pendiente():
            available_notes = cargar_notas_de_disco()
            if available_notes:
                st.session_state.partial_notes = available_notes
                st.session_state.progreso_recuperado_disco = True
        
        if not available_notes:
            st.info("👆 Genera las notas primero en el Paso 2")
            
            # v2.9.0: Verificar si hay algo en disco
            if hay_progreso_pendiente():
                resumen = obtener_resumen_progreso()
                if resumen and resumen['notas'] > 0:
                    st.warning(f"💾 Se detectaron **{resumen['notas']} notas** guardadas en disco.")
                    if st.button("📂 Recuperar notas de disco"):
                        notas = cargar_notas_de_disco()
                        if notas:
                            st.session_state.partial_notes = notas
                            st.session_state.progreso_recuperado_disco = True
                            st.rerun()
            return
        
        generated_notes = available_notes
        
        # v2.9.0: Indicador de origen
        if st.session_state.get('progreso_recuperado_disco'):
            st.success(f"✅ **{len(generated_notes)} notas listas para descargar** (💾 recuperadas de disco)")
        else:
            st.success(f"✅ **{len(generated_notes)} notas listas para descargar**")
        
        # Estadísticas
        ajustadas = st.session_state.get('notas_ajustadas', [])
        if ajustadas:
            st.info(f"⚙️ {len(ajustadas)} notas fueron ajustadas automáticamente")
        
        # Crear ZIP
        zip_buffer = crear_zip_parcial(generated_notes)
        
        # Nombre del archivo
        fecha_str = st.session_state.get('fecha_nota', date.today())
        if fecha_str:
            fecha_archivo = fecha_str.strftime('%Y%m%d')
        else:
            fecha_archivo = datetime.now().strftime('%Y%m%d_%H%M%S')
        
        col_download1, col_download2 = st.columns(2)
        
        with col_download1:
            st.download_button(
                label=f"📦 Descargar todas ({len(generated_notes)} notas) - ZIP",
                data=zip_buffer,
                file_name=f"notas_nefrologicas_{fecha_archivo}.zip",
                mime="application/zip",
                use_container_width=True
            )
        
        with col_download2:
            if st.button("🗑️ Limpiar y empezar de nuevo", use_container_width=True):
                limpiar_estado_generacion()
                st.rerun()
        
        st.markdown("---")
        
        # Vista previa
        st.subheader("👁️ Vista previa de notas")
        
        opciones = [f"{n['expediente']} - {n['paciente']}" + (" ⚙️" if n.get('fue_ajustada') else "") for n in generated_notes]
        
        selected_note = st.selectbox(
            "Selecciona una nota para ver",
            options=opciones,
            help="⚙️ = nota ajustada automáticamente"
        )
        
        if selected_note:
            selected_clean = selected_note.replace(" ⚙️", "")
            opciones_clean = [o.replace(" ⚙️", "") for o in opciones]
            note_idx = opciones_clean.index(selected_clean)
            note = generated_notes[note_idx]
            
            col1, col2 = st.columns([3, 1])
            
            with col1:
                nota_len = len(note['content'])
                status_text_display = f"📊 {nota_len} caracteres"
                if note.get('fue_ajustada'):
                    status_text_display += " (⚙️ ajustada automáticamente)"
                st.caption(status_text_display)
                
                st.text_area(
                    "Contenido de la nota",
                    value=note['content'],
                    height=400
                )
            
            with col2:
                st.download_button(
                    label="⬇️ Descargar esta nota",
                    data=note['doc_bytes'],
                    file_name=f"{note['expediente']}_{note['paciente'].replace(' ', '_')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
        
        # v2.9.0: Información sobre almacenamiento en disco
        st.markdown("---")
        st.subheader("💾 Almacenamiento en Disco")
        
        if hay_progreso_pendiente():
            resumen = obtener_resumen_progreso()
            if resumen:
                col_d1, col_d2, col_d3 = st.columns(3)
                with col_d1:
                    st.metric("📄 Notas en disco", resumen['notas'])
                with col_d2:
                    st.metric("✅ Expedientes completados", resumen['completados'])
                with col_d3:
                    st.metric("❌ Errores", resumen['errores'])
                
                if resumen['ultima_actualizacion']:
                    st.caption(f"Última actualización: {resumen['ultima_actualizacion']}")
        else:
            st.info("📭 No hay progreso guardado en disco actualmente.")

    # ========================================================================
    # v2.10.0: TAB 4: GESTIÓN DE FECHAS DE NACIMIENTO
    # ========================================================================
    with tab4:
        st.markdown('<div class="step-header"><h2>📅 Gestión de Fechas de Nacimiento</h2></div>', unsafe_allow_html=True)
        
        st.markdown("""
        El sistema guarda las **fechas de nacimiento (FN)** de cada paciente en un caché persistente.
        Esto permite calcular automáticamente la edad correcta en cada nota mensual, considerando 
        si el paciente ya cumplió años en el año en curso.
        
        **Fuentes de FN (en orden de prioridad):**
        1. 💾 Caché persistente (este archivo)
        2. 📄 Extracción automática de notas HD previas (auto-guarda al caché)
        3. ✍️ Ingreso manual aquí
        """)
        
        cache_fn = cargar_cache_fn()
        
        col_stat1, col_stat2, col_stat3 = st.columns(3)
        with col_stat1:
            st.metric("📅 Pacientes en caché", len(cache_fn))
        
        # Si hay censo cargado, mostrar cuáles faltan
        censo_fn = st.session_state.get('censo_df', None)
        if censo_fn is not None:
            expedientes_censo = set(censo_fn['Exped.2'].astype(str).str.upper().str.strip())
            expedientes_cache = set(cache_fn.keys())
            faltantes = expedientes_censo - expedientes_cache
            
            with col_stat2:
                st.metric("✅ Censo con FN", len(expedientes_censo & expedientes_cache))
            with col_stat3:
                st.metric("⚠️ Censo sin FN", len(faltantes))
        
        st.markdown("---")
        
        # ====================================================================
        # SECCIÓN 1: AGREGAR/EDITAR FN INDIVIDUAL
        # ====================================================================
        st.subheader("✍️ Agregar o editar FN")
        
        col_form1, col_form2, col_form3 = st.columns([2, 2, 1])
        
        with col_form1:
            if censo_fn is not None:
                # Selector con expedientes del censo
                opciones_exp = sorted(censo_fn['Exped.2'].astype(str).str.upper().str.strip().tolist())
                exp_input = st.selectbox(
                    "Expediente",
                    options=[""] + opciones_exp,
                    format_func=lambda x: f"{x}{' ✅' if x in cache_fn else ' ⚠️'}" if x else "— Selecciona —",
                    key="fn_exp_select"
                )
            else:
                exp_input = st.text_input(
                    "Expediente",
                    placeholder="Ej: RL0069",
                    key="fn_exp_text"
                ).upper().strip()
        
        with col_form2:
            # Si el expediente ya está en caché, pre-cargar el valor
            valor_actual = cache_fn.get(exp_input, "") if exp_input else ""
            
            # Key dinámica para que el campo se refresque al cambiar de expediente
            fn_input = st.text_input(
                "Fecha de nacimiento (dd/mm/aaaa)",
                value=valor_actual,
                placeholder="04/05/1955",
                key=f"fn_fecha_input_{exp_input or 'vacio'}"
            )
        
        with col_form3:
            st.write("")  # Spacer
            st.write("")  # Spacer
            guardar_btn = st.button("💾 Guardar", use_container_width=True, key="fn_guardar_btn")
        
        if guardar_btn:
            if not exp_input:
                st.error("❌ Selecciona o ingresa un expediente")
            elif not fn_input:
                st.error("❌ Ingresa una fecha de nacimiento")
            else:
                # Validar formato
                try:
                    fn_obj = datetime.strptime(fn_input, '%d/%m/%Y')
                    if fn_obj.year < 1900 or fn_obj > datetime.now():
                        st.error("❌ Fecha de nacimiento fuera de rango razonable")
                    else:
                        if agregar_fn_a_cache(exp_input, fn_input):
                            edad_calc = calcular_edad(fn_input, date.today())
                            st.success(f"✅ Guardado: **{exp_input}** → FN **{fn_input}** (edad actual: {edad_calc} años)")
                            st.rerun()
                        else:
                            st.error("❌ Error al guardar")
                except ValueError:
                    st.error("❌ Formato inválido. Usa dd/mm/aaaa (ej: 04/05/1955)")
        
        # Mostrar edad calculada en vivo si hay FN válida
        if fn_input and exp_input:
            try:
                fn_obj_preview = datetime.strptime(fn_input, '%d/%m/%Y')
                fecha_nota_actual = st.session_state.get('fecha_nota', None) or date.today()
                edad_preview = calcular_edad(fn_input, fecha_nota_actual)
                if edad_preview is not None:
                    st.info(f"👁️ Edad calculada a la fecha de la nota ({fecha_nota_actual if isinstance(fecha_nota_actual, date) else fecha_nota_actual.strftime('%d/%m/%Y')}): **{edad_preview} años**")
            except ValueError:
                pass
        
        st.markdown("---")
        
        # ====================================================================
        # SECCIÓN 2: PACIENTES SIN FN (DEL CENSO)
        # ====================================================================
        if censo_fn is not None and faltantes:
            st.subheader(f"⚠️ Pacientes del censo sin FN ({len(faltantes)})")
            
            with st.expander("Ver lista de pacientes faltantes", expanded=False):
                # Crear DataFrame con info de los faltantes
                faltantes_data = []
                for exp in sorted(faltantes):
                    nombre_row = censo_fn[censo_fn['Exped.2'].astype(str).str.upper().str.strip() == exp]
                    nombre = nombre_row.iloc[0]['Paciente'] if not nombre_row.empty else "—"
                    faltantes_data.append({
                        'Expediente': exp,
                        'Paciente': nombre
                    })
                
                df_faltantes = pd.DataFrame(faltantes_data)
                st.dataframe(df_faltantes, use_container_width=True, height=300)
                
                # Descargar plantilla
                csv_plantilla = df_faltantes.copy()
                csv_plantilla['FN'] = ''
                csv_buffer = io.StringIO()
                csv_plantilla.to_csv(csv_buffer, index=False)
                st.download_button(
                    "📥 Descargar plantilla CSV para llenar",
                    data=csv_buffer.getvalue(),
                    file_name="pacientes_sin_fn.csv",
                    mime="text/csv"
                )
            
            # Carga masiva por CSV
            st.markdown("**Carga masiva desde CSV:**")
            st.caption("Sube un CSV con columnas `Expediente` y `FN` (formato dd/mm/aaaa)")
            
            csv_fn_file = st.file_uploader("CSV con FNs", type=['csv'], key='csv_fn_upload')
            if csv_fn_file:
                try:
                    df_fn_carga = pd.read_csv(csv_fn_file)
                    cols_norm = {c.lower().strip(): c for c in df_fn_carga.columns}
                    
                    col_exp = cols_norm.get('expediente') or cols_norm.get('exp')
                    col_fn = cols_norm.get('fn') or cols_norm.get('fecha de nacimiento') or cols_norm.get('fecha_nacimiento')
                    
                    if not col_exp or not col_fn:
                        st.error("❌ El CSV debe tener columnas 'Expediente' y 'FN'")
                    else:
                        if st.button("⬆️ Procesar CSV", key="procesar_csv_fn"):
                            agregados = 0
                            errores_csv = []
                            for _, row in df_fn_carga.iterrows():
                                exp = str(row[col_exp]).strip().upper()
                                fn_val = str(row[col_fn]).strip()
                                if not exp or not fn_val or fn_val.lower() in ('nan', 'none', ''):
                                    continue
                                if agregar_fn_a_cache(exp, fn_val):
                                    agregados += 1
                                else:
                                    errores_csv.append(f"{exp}: '{fn_val}'")
                            
                            st.success(f"✅ {agregados} FNs agregadas al caché")
                            if errores_csv:
                                with st.expander(f"⚠️ {len(errores_csv)} errores"):
                                    for err in errores_csv:
                                        st.text(err)
                            time.sleep(1)
                            st.rerun()
                except Exception as e:
                    st.error(f"❌ Error procesando CSV: {e}")
            
            st.markdown("---")
        
        # ====================================================================
        # SECCIÓN 3: TABLA COMPLETA DEL CACHÉ
        # ====================================================================
        st.subheader(f"💾 Caché completo ({len(cache_fn)} pacientes)")
        
        if cache_fn:
            # Construir DataFrame para mostrar
            cache_data = []
            for exp, fn_val in sorted(cache_fn.items()):
                edad_actual = calcular_edad(fn_val, date.today())
                cache_data.append({
                    'Expediente': exp,
                    'Fecha de nacimiento': fn_val,
                    'Edad actual': f"{edad_actual} años" if edad_actual is not None else "—"
                })
            
            df_cache = pd.DataFrame(cache_data)
            st.dataframe(df_cache, use_container_width=True, height=400)
            
            # Descarga del caché completo
            col_dl1, col_dl2 = st.columns(2)
            with col_dl1:
                csv_cache_buf = io.StringIO()
                df_cache.to_csv(csv_cache_buf, index=False)
                st.download_button(
                    "📥 Descargar caché como CSV",
                    data=csv_cache_buf.getvalue(),
                    file_name=f"cache_fn_{datetime.now().strftime('%Y%m%d')}.csv",
                    mime="text/csv"
                )
            
            with col_dl2:
                json_cache_buf = json.dumps({'pacientes': cache_fn}, ensure_ascii=False, indent=2)
                st.download_button(
                    "📥 Descargar caché como JSON",
                    data=json_cache_buf,
                    file_name=f"cache_fn_{datetime.now().strftime('%Y%m%d')}.json",
                    mime="application/json"
                )
            
            st.markdown("---")
            
            # Eliminar entradas
            st.markdown("**🗑️ Eliminar entrada del caché:**")
            col_del1, col_del2 = st.columns([3, 1])
            with col_del1:
                exp_eliminar = st.selectbox(
                    "Expediente a eliminar",
                    options=[""] + sorted(cache_fn.keys()),
                    key="fn_eliminar_select"
                )
            with col_del2:
                st.write("")
                st.write("")
                if st.button("🗑️ Eliminar", use_container_width=True, key="fn_eliminar_btn"):
                    if exp_eliminar and exp_eliminar in cache_fn:
                        del cache_fn[exp_eliminar]
                        if guardar_cache_fn(cache_fn):
                            st.success(f"✅ {exp_eliminar} eliminado del caché")
                            time.sleep(1)
                            st.rerun()
        else:
            st.info("📭 El caché está vacío. Las FNs se irán agregando automáticamente cuando se procesen notas HD que las contengan, o las puedes agregar manualmente arriba.")

# ============================================================================
# EJECUTAR APP
# ============================================================================

if __name__ == "__main__":
    main()
